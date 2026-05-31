"""
本科毕业设计(论文) Word 格式检测脚本

用法：
  python thesis_format_audit.py "论文.docx"
  python thesis_format_audit.py "论文.docx" --out report.html

说明：
  - 只读取 docx，不修改原文件。
  - 按《附件3：本科毕业设计(论文)参考模板》中的常见格式要求做结构化检测。
  - 输出控制台摘要，并生成 HTML 报告。
"""
from __future__ import annotations

import argparse
import ast
import hashlib
import html
import json
import re
import sys
import tempfile
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("缺少 python-docx。请先安装：pip install python-docx", file=sys.stderr)
    raise

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
W_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
OFFICIAL_FIXED_HEADER = "中国石油大学（华东）本科毕业设计（论文）"
LEGACY_FIXED_HEADER = "中国石油大学（华东）本科毕业设计(论文)"
STRICT_REL_PREFIX = "http://purl.oclc.org/ooxml/officeDocument/relationships"
TRANSITIONAL_REL_PREFIX = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
STRICT_TO_TRANSITIONAL_URIS = {
    STRICT_REL_PREFIX: TRANSITIONAL_REL_PREFIX,
    "http://purl.oclc.org/ooxml/wordprocessingml/main": W_URI,
    "http://purl.oclc.org/ooxml/officeDocument/math": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "http://purl.oclc.org/ooxml/drawingml/main": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "http://purl.oclc.org/ooxml/drawingml/picture": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "http://purl.oclc.org/ooxml/officeDocument/sharedTypes": "http://schemas.openxmlformats.org/officeDocument/2006/sharedTypes",
    "http://purl.oclc.org/ooxml/officeDocument/extendedProperties": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    "http://purl.oclc.org/ooxml/officeDocument/customProperties": "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties",
}


def qn(tag: str) -> str:
    return W_NS + tag


@dataclass
class Finding:
    item: str
    status: str  # PASS / FAIL / WARN
    location: str
    message: str
    evidence: str = ""
    severity: str = "一般"


@dataclass
class AuditContext:
    path: Path
    doc: Document
    findings: list[Finding] = field(default_factory=list)
    _paragraph_texts: list[str] = field(default_factory=list, init=False, repr=False)
    _paragraph_clean_texts: list[str] = field(default_factory=list, init=False, repr=False)
    _paragraph_text_join: Optional[str] = field(default=None, init=False, repr=False)
    _body_blocks: Optional[list[tuple[str, object]]] = field(default=None, init=False, repr=False)
    _document_xml: Optional[str] = field(default=None, init=False, repr=False)
    _media_count: Optional[int] = field(default=None, init=False, repr=False)
    _first_chapter_index: Optional[int] = field(default=None, init=False, repr=False)
    _first_chapter_index_ready: bool = field(default=False, init=False, repr=False)
    _exact_index_cache: dict[str, Optional[int]] = field(default_factory=dict, init=False, repr=False)

    def add(self, item: str, status: str, location: str, message: str, evidence: str = "", severity: str = "一般"):
        self.findings.append(Finding(item, status, location, message, evidence, severity))

    @property
    def paragraph_texts(self) -> list[str]:
        if not self._paragraph_texts:
            self._paragraph_texts = [p.text for p in self.doc.paragraphs]
        return self._paragraph_texts

    @property
    def paragraph_clean_texts(self) -> list[str]:
        if not self._paragraph_clean_texts:
            self._paragraph_clean_texts = [clean_text(text) for text in self.paragraph_texts]
        return self._paragraph_clean_texts

    @property
    def full_text(self) -> str:
        if self._paragraph_text_join is None:
            self._paragraph_text_join = "\n".join(self.paragraph_texts)
        return self._paragraph_text_join

    def text_until(self, end: Optional[int]) -> str:
        if end is None:
            return self.full_text
        return "\n".join(self.paragraph_texts[:end])

    @property
    def body_blocks(self) -> list[tuple[str, object]]:
        if self._body_blocks is None:
            self._body_blocks = list(iter_body_blocks(self.doc))
        return self._body_blocks

    @property
    def document_xml(self) -> str:
        if self._document_xml is None:
            self._document_xml = document_xml(self.path)
        return self._document_xml

    @property
    def media_count(self) -> int:
        if self._media_count is None:
            self._media_count = media_count(self.path)
        return self._media_count

    def find_exact(self, exact: str) -> Optional[int]:
        if exact not in self._exact_index_cache:
            self._exact_index_cache[exact] = find_paragraph_index(self.doc, exact)
        return self._exact_index_cache[exact]

    @property
    def first_chapter_index(self) -> Optional[int]:
        if not self._first_chapter_index_ready:
            self._first_chapter_index = find_first_chapter_index(self.doc)
            self._first_chapter_index_ready = True
        return self._first_chapter_index

    def first_tail_index(self) -> int:
        candidates = [
            i for i, text in enumerate(self.paragraph_clean_texts)
            if text in ("致谢", "参考文献", "附录")
        ]
        return min(candidates) if candidates else len(self.doc.paragraphs)


def clean_text(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def run_east_asia(run) -> Optional[str]:
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("eastAsia"))
    return None


def run_size(run) -> Optional[float]:
    return run.font.size.pt if run.font.size is not None else None


def has_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def has_latin_or_digit(text: str) -> bool:
    return bool(re.search(r"[A-Za-z0-9\u00b2\u00b3\u00b9\u2070\u2074-\u2079\u2080-\u2089]", text or ""))


SUPERSCRIPT_DIGITS = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
SUBSCRIPT_DIGITS = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")


def math_script_display(text: str, kind: str) -> str:
    if kind == "sub":
        return text.translate(SUBSCRIPT_DIGITS)
    return text.translate(SUPERSCRIPT_DIGITS)


def chapter_number_to_chinese(number: int) -> str:
    digits = "零一二三四五六七八九"
    if number <= 0:
        return str(number)
    if number < 10:
        return digits[number]
    if number == 10:
        return "十"
    if number < 20:
        return "十" + digits[number % 10]
    if number < 100:
        tens, ones = divmod(number, 10)
        return digits[tens] + "十" + (digits[ones] if ones else "")
    return str(number)


def short_sample(text: str, limit: int = 46) -> str:
    clean = re.sub(r"\s+", " ", text or "").strip()
    return clean if len(clean) <= limit else clean[:limit] + "..."


def first_text_run(paragraph):
    for r in paragraph.runs:
        if r.text and r.text.strip():
            return r
    return None


def style_east_asia(style) -> Optional[str]:
    rpr = getattr(style._element, "rPr", None)
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("eastAsia"))
    return None


def style_size(style) -> Optional[float]:
    return style.font.size.pt if style.font.size is not None else None


def iter_style_chain(style):
    """Yield the style and its base styles, closest style first."""
    seen = set()
    cur = style
    while cur is not None and id(cur) not in seen:
        seen.add(id(cur))
        yield cur
        cur = getattr(cur, "base_style", None)


def effective_run_size(paragraph) -> Optional[float]:
    run = first_text_run(paragraph)
    direct = run_size(run) if run else None
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = style_size(style)
        if value is not None:
            return value
    return None


def effective_run_east_asia(paragraph) -> Optional[str]:
    run = first_text_run(paragraph)
    direct = run_east_asia(run) if run else None
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = style_east_asia(style)
        if value is not None:
            return value
    return None


def effective_run_latin_font(paragraph) -> Optional[str]:
    run = first_text_run(paragraph)
    if run is not None:
        direct = run_latin_font(run)
        if direct:
            return direct
    for style in iter_style_chain(paragraph.style):
        if style.font.name:
            return style.font.name
    return None


def effective_size_for_run(paragraph, run) -> Optional[float]:
    direct = run_size(run)
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = style_size(style)
        if value is not None:
            return value
    return None


def effective_east_asia_for_run(paragraph, run) -> Optional[str]:
    direct = run_east_asia(run)
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = style_east_asia(style)
        if value is not None:
            return value
    return None


def effective_latin_for_run(paragraph, run) -> Optional[str]:
    direct = run_latin_font(run)
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        if style.font.name:
            return style.font.name
    return None


def run_latin_font(run) -> Optional[str]:
    if run.font.name:
        return run.font.name
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("ascii")) or rpr.rFonts.get(qn("hAnsi"))
    return None


def is_kaiti_font(name: Optional[str]) -> bool:
    return name in (None, "楷体_GB2312", "楷体", "KaiTi", "KaiTi_GB2312")


def effective_run_bold(paragraph) -> Optional[bool]:
    run = first_text_run(paragraph)
    if run is not None and run.font.bold is not None:
        return run.font.bold
    for style in iter_style_chain(paragraph.style):
        if style.font.bold is not None:
            return style.font.bold
    return None


def para_line_spacing(paragraph):
    return paragraph.paragraph_format.line_spacing


def effective_line_spacing(paragraph):
    direct = paragraph.paragraph_format.line_spacing
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = style.paragraph_format.line_spacing
        if value is not None:
            return value
    return None


def ppr_spacing_attr(ppr, attr: str) -> Optional[str]:
    if ppr is None:
        return None
    spacing = getattr(ppr, "spacing", None)
    if spacing is None:
        return None
    return spacing.get(qn(attr))


def para_has_half_line_spacing(paragraph, attr: str) -> bool:
    spacing = paragraph._p.pPr.spacing if paragraph._p.pPr is not None else None
    if spacing is None:
        return False
    value = spacing.get(qn(attr))
    return value == "50"


def effective_half_line_spacing(paragraph, attr: str) -> bool:
    direct = ppr_spacing_attr(paragraph._p.pPr, attr)
    if direct is not None:
        return direct == "50"
    for style in iter_style_chain(paragraph.style):
        ppr = getattr(style._element, "pPr", None)
        value = ppr_spacing_attr(ppr, attr)
        if value is not None:
            return value == "50"
    return False


def para_has_first_line_chars(paragraph, chars: int = 200) -> bool:
    ppr = paragraph._p.pPr
    if ppr is None or ppr.ind is None:
        return False
    return ppr.ind.get(qn("firstLineChars")) == str(chars)


def effective_first_line_chars(paragraph, chars: int = 200) -> bool:
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.ind is not None:
        value = ppr.ind.get(qn("firstLineChars"))
        if value is not None:
            return value == str(chars)
    for style in iter_style_chain(paragraph.style):
        ppr = getattr(style._element, "pPr", None)
        if ppr is not None and ppr.ind is not None:
            value = ppr.ind.get(qn("firstLineChars"))
            if value is not None:
                return value == str(chars)
    return False


def effective_first_line_indent_ok(paragraph) -> bool:
    if effective_first_line_chars(paragraph, 200):
        return True
    ppr = paragraph._p.pPr
    candidates = []
    if ppr is not None and ppr.ind is not None:
        candidates.append(ppr.ind)
    for style in iter_style_chain(paragraph.style):
        sppr = getattr(style._element, "pPr", None)
        if sppr is not None and sppr.ind is not None:
            candidates.append(sppr.ind)
    for ind in candidates:
        first_line = ind.get(qn("firstLine"))
        if first_line is None:
            continue
        try:
            twips = int(first_line)
        except ValueError:
            continue
        if 400 <= twips <= 560:
            return True
    return False


def is_subfigure_label_paragraph(text: str) -> bool:
    if not text:
        return False
    if not re.match(r"^\s*[（(][a-z][）)]", text, flags=re.IGNORECASE):
        return False
    return len(text) < 140 and not re.search(r"[。；;，,]$", text)


def para_has_snap_to_grid(paragraph) -> bool:
    ppr = paragraph._p.pPr
    if ppr is None:
        return False
    snap = ppr.find(qn("snapToGrid"))
    if snap is None:
        return False
    value = snap.get(qn("val"))
    return value in (None, "1", "true", "on")


def ppr_snap_to_grid_value(ppr) -> Optional[bool]:
    if ppr is None:
        return None
    snap = ppr.find(qn("snapToGrid"))
    if snap is None:
        return None
    value = snap.get(qn("val"))
    return value in (None, "1", "true", "on")


def effective_snap_to_grid(paragraph) -> bool:
    direct = ppr_snap_to_grid_value(paragraph._p.pPr)
    if direct is not None:
        return direct
    for style in iter_style_chain(paragraph.style):
        value = ppr_snap_to_grid_value(getattr(style._element, "pPr", None))
        if value is not None:
            return value
    return False


def effective_paragraph_alignment(paragraph):
    if paragraph.alignment is not None:
        return paragraph.alignment
    for style in iter_style_chain(paragraph.style):
        value = style.paragraph_format.alignment
        if value is not None:
            return value
    return None


def paragraph_font_issues(paragraph, expected_size: float = 12, expected_east: str = "宋体", expected_latin: str = "Times New Roman"):
    issues = []
    for r in paragraph.runs:
        if not r.text:
            continue
        if re.fullmatch(r"\[\d+(?:\s*[-,，]\s*\d+)*\]", r.text.strip()):
            continue
        size = effective_size_for_run(paragraph, r)
        east = effective_east_asia_for_run(paragraph, r)
        latin = effective_latin_for_run(paragraph, r)
        if size is not None and abs(size - expected_size) > 0.2:
            issues.append(f"字号{size}")
        if has_cjk(r.text) and east is not None and east != expected_east:
            issues.append(f"中文字体{east}")
        if has_latin_or_digit(r.text) and latin is not None and latin != expected_latin:
            issues.append(f"西文字体{latin}")
    return issues


def caption_run_segments(paragraph, text: str, num: str, title: str) -> list[tuple[str, str, list]]:
    """Group a visible caption into Word-like fragments for human-readable checks."""
    parts = [(r.text, r) for r in paragraph.runs if r.text]
    prefix_m = re.match(r"^(图\d+-\d+)\s*(.*?)(\[\d+(?:[-,，]\d+)*\])?$", text.strip())
    if prefix_m:
        label = prefix_m.group(1)
        title_text = (prefix_m.group(2) or "").strip()
        cite = prefix_m.group(3) or ""
    else:
        label = num
        title_text = title.strip()
        cite_m = re.search(r"(\[\d+(?:[-,，]\d+)*\])\s*$", title_text)
        cite = cite_m.group(1) if cite_m else ""
        if cite:
            title_text = title_text[: -len(cite)].strip()

    def collect(target: str) -> list:
        if not target:
            return []
        compact_target = re.sub(r"\s+", "", target)
        collected = ""
        runs = []
        started = False
        for run_text, run in parts:
            compact = re.sub(r"\s+", "", run_text)
            if not compact:
                continue
            candidate = collected + compact
            if not started:
                if compact_target.startswith(compact):
                    started = True
                elif compact in compact_target or compact_target.startswith(candidate):
                    started = True
                else:
                    continue
            runs.append(run)
            collected = candidate
            if collected == compact_target or compact_target.startswith(collected):
                if collected == compact_target:
                    break
            elif collected.endswith(compact_target) or compact_target in collected:
                break
        return runs

    segments = []
    if label:
        segments.append(("图号标签", label, collect(label)))
    if title_text:
        segments.append(("图题文字", title_text, collect(title_text)))
    if cite:
        segments.append(("文献标注", cite, collect(cite)))
    return segments


def caption_segment_issues(paragraph, segment_text: str, runs: list, expected_size: float = 12, strict_size: bool = True) -> list[str]:
    issues = []
    size_values = []
    east_values = []
    latin_values = []
    target_runs = runs or [r for r in paragraph.runs if r.text and r.text.strip()]
    for run in target_runs:
        if not run.text or not run.text.strip():
            continue
        r_size = effective_size_for_run(paragraph, run)
        r_east = effective_east_asia_for_run(paragraph, run)
        r_latin = effective_latin_for_run(paragraph, run)
        if r_size is not None and r_size not in size_values:
            size_values.append(r_size)
        if has_cjk(run.text) and r_east is not None and r_east not in east_values:
            east_values.append(r_east)
        if has_latin_or_digit(run.text) and r_latin is not None and r_latin not in latin_values:
            latin_values.append(r_latin)

    bad_sizes = [s for s in size_values if abs(s - expected_size) > 0.2]
    if not strict_size:
        bad_sizes = [s for s in bad_sizes if not any(abs(s - ok) <= 0.2 for ok in (10.5, 12))]
    if bad_sizes:
        shown = "/".join(f"{s:g}" for s in bad_sizes[:3])
        issues.append(f"当前字号{shown}pt，应为小四12pt")
    bad_east = [f for f in east_values if f != "宋体"]
    if has_cjk(segment_text) and bad_east:
        issues.append(f"当前中文字体{'/'.join(bad_east[:3])}，应为宋体")
        near_five = [s for s in size_values if abs(s - 10.5) <= 0.2]
        if not strict_size and near_five:
            issues.append("当前字号五号，按维普图题文字规范应为小四12pt")
    bad_latin = [f for f in latin_values if f != "Times New Roman"]
    if has_latin_or_digit(segment_text) and bad_latin:
        issues.append(f"当前西文字体{'/'.join(bad_latin[:3])}，应为Times New Roman")
    return issues


def vip_figure_label_issues(paragraph, segment_text: str, runs: list) -> list[str]:
    issues = []
    target_runs = runs or [r for r in paragraph.runs if r.text and r.text.strip()]
    size_values = []
    east_values = []
    for run in target_runs:
        if not run.text or not run.text.strip():
            continue
        r_size = effective_size_for_run(paragraph, run)
        r_east = effective_east_asia_for_run(paragraph, run)
        if r_size is not None and r_size not in size_values:
            size_values.append(r_size)
        if has_cjk(run.text) and r_east is not None and r_east not in east_values:
            east_values.append(r_east)
    bad_sizes = [s for s in size_values if abs(s - 10.5) > 0.2]
    if bad_sizes:
        issues.append(f"当前字号{'/'.join(f'{s:g}' for s in bad_sizes[:3])}pt，应为五号10.5pt")
    bad_east = [f for f in east_values if f != "宋体"]
    if has_cjk(segment_text) and bad_east:
        issues.append(f"当前中文字体{'/'.join(bad_east[:3])}，应为宋体")
    return issues


def leading_label_bold(paragraph, label: str) -> bool:
    collected = ""
    label_runs = []
    for r in paragraph.runs:
        if not r.text:
            continue
        need = len(label) - len(collected)
        if need <= 0:
            break
        collected += r.text[:need]
        label_runs.append(r)
        if len(collected) >= len(label):
            break
    return collected == label and bool(label_runs) and all((r.font.bold is True or r.bold is True) for r in label_runs)


def is_probable_body_paragraph(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or has_drawing(paragraph):
        return False
    style = paragraph.style.name
    if style.startswith("Heading") or style.startswith("toc") or style in ("Caption",):
        return False
    if re.match(r"^(图|表)\d+[-－]\d+", text):
        return False
    if re.match(r"^\d+\.\d+(?:\.\d+)?\s+\S+", text):
        return False
    if re.match(r"^\[\d+\]", text):
        return False
    if text in ("摘  要", "Abstract", "目  录", "致  谢", "参考文献", "附  录"):
        return False
    if text.startswith(("关键词", "Keywords")):
        return False
    if is_subfigure_label_paragraph(text):
        return False
    if len(text) < 25:
        return False
    return True


def is_body_or_caption_paragraph(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or has_drawing(paragraph):
        return False
    style = paragraph.style.name
    if style.startswith(("Heading", "toc", "目录", "Bibliography")):
        return False
    if re.match(r"^\[\d+\]", text):
        return False
    if text in ("摘  要", "Abstract", "目  录", "致  谢", "参考文献", "附  录"):
        return False
    if text.startswith(("关键词", "Keywords")):
        return False
    return True


def chapter_number_from_text(text: str) -> Optional[int]:
    m = re.match(r"^第\s*(\d+)\s*章", text.strip())
    if m:
        return int(m.group(1))
    cn = re.match(r"^第\s*([一二三四五六七八九十]+)\s*章", text.strip())
    if not cn:
        return None
    digits = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    value = cn.group(1)
    if value == "十":
        return 10
    if value.startswith("十"):
        return 10 + digits.get(value[-1], 0)
    if "十" in value:
        left, right = value.split("十", 1)
        return digits.get(left, 0) * 10 + (digits.get(right, 0) if right else 0)
    return digits.get(value)


def paragraph_has_numbering(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.find(qn("numPr")) is not None


def paragraph_num_id(paragraph) -> Optional[str]:
    ppr = paragraph._p.pPr
    num_pr = ppr.find(qn("numPr")) if ppr is not None else None
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("numId"))
    return num_id.get(qn("val")) if num_id is not None else None


def numbering_formats(path: Path) -> dict[str, str]:
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/numbering.xml")
    except Exception:
        return {}
    ns = {"w": W_URI}
    root = ET.fromstring(xml)
    abstract_text: dict[str, str] = {}
    for absn in root.findall("w:abstractNum", ns):
        abs_id = absn.get(qn("abstractNumId"))
        lvl = absn.find("w:lvl[@w:ilvl='0']", ns)
        if lvl is None:
            lvl = absn.find("w:lvl", ns)
        if abs_id is None or lvl is None:
            continue
        text = lvl.find("w:lvlText", ns)
        if text is not None:
            abstract_text[abs_id] = text.get(qn("val"), "")
    result = {}
    for num in root.findall("w:num", ns):
        num_id = num.get(qn("numId"))
        abs_el = num.find("w:abstractNumId", ns)
        abs_id = abs_el.get(qn("val")) if abs_el is not None else None
        if num_id and abs_id in abstract_text:
            result[num_id] = abstract_text[abs_id]
    return result


def has_drawing(paragraph) -> bool:
    xml = paragraph._p.xml
    return "a:blip" in xml or "<v:imagedata" in xml or "imagedata" in xml


def is_centered(paragraph) -> bool:
    return effective_paragraph_alignment(paragraph) == WD_ALIGN_PARAGRAPH.CENTER


def iter_body_blocks(doc: Document):
    for child in doc._element.body.iterchildren():
        if child.tag == qn("p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("tbl"):
            yield "t", Table(child, doc)


def table_plain_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def is_formula_table(table: Table) -> bool:
    if len(table.rows) != 1 or len(table.columns) != 2:
        return False
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    has_math = any(p._p.xpath(".//m:oMath") or p._p.xpath(".//m:oMathPara") for p in left.paragraphs)
    num = " ".join(p.text.strip() for p in right.paragraphs if p.text.strip())
    return has_math or bool(re.fullmatch(r"\((?:\d+-\d+|A-\d+)\)", num))


def is_cover_table(table: Table) -> bool:
    text = table_plain_text(table)
    return "题  目" in text or "学生姓名" in text or "指导教师" in text


def find_paragraph_index(doc: Document, exact: str) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == exact:
            return i
    return None


def find_first(doc: Document, predicate) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    return None


def media_count(path: Path) -> int:
    try:
        with zipfile.ZipFile(path) as z:
            return len([n for n in z.namelist() if n.startswith("word/media/")])
    except Exception:
        return 0


def document_xml(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as z:
            return z.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return ""


def docx_has_toc_field(path: Path) -> bool:
    xml = document_xml(path)
    return "TOC \\o" in xml or "TOC" in xml and "PAGEREF" in xml


def xml_has_toc_field(xml: str) -> bool:
    return "TOC \\o" in xml or ("TOC" in xml and "PAGEREF" in xml)


def docx_pageref_count(path: Path) -> int:
    return document_xml(path).count("PAGEREF")


def make_python_docx_compatible_copy(path: Path, work_dir: Path) -> Path:
    """Convert Strict OOXML relationship types in a temp copy for python-docx."""
    try:
        with zipfile.ZipFile(path) as src:
            names = src.namelist()
            xml_names = [
                name for name in names
                if name.endswith((".xml", ".rels")) and not name.startswith("word/media/")
            ]
            needs_rewrite = any(
                any(strict in src.read(name).decode("utf-8", errors="ignore") for strict in STRICT_TO_TRANSITIONAL_URIS)
                for name in xml_names
            )
            if not needs_rewrite:
                return path

            fixed = work_dir / f"{path.stem}_compatible.docx"
            with zipfile.ZipFile(fixed, "w", compression=zipfile.ZIP_DEFLATED) as dst:
                for info in src.infolist():
                    data = src.read(info.filename)
                    if info.filename in xml_names:
                        for strict, transitional in STRICT_TO_TRANSITIONAL_URIS.items():
                            data = data.replace(strict.encode("utf-8"), transitional.encode("utf-8"))
                    dst.writestr(info, data)
            return fixed
    except zipfile.BadZipFile as exc:
        raise ValueError("这个文件扩展名是 .docx，但内部不是有效的 Word 文档包。请在 Word/WPS 中另存为 .docx 后再上传。") from exc


def open_docx_document(path: Path, work_dir: Optional[Path] = None) -> tuple[Document, Path]:
    if work_dir is None:
        with tempfile.TemporaryDirectory(prefix="docx-compat-") as tmp:
            return open_docx_document(path, Path(tmp))

    compatible_path = make_python_docx_compatible_copy(path, work_dir)
    try:
        return Document(compatible_path), compatible_path
    except KeyError as exc:
        if "officeDocument" in str(exc):
            raise ValueError(
                "这个 DOCX 的内部关系文件不标准，无法定位正文 document.xml。"
                "请在 Word 或 WPS 中打开后选择“另存为 .docx”，再重新上传。"
            ) from exc
        raise
    except Exception as exc:
        raise ValueError(f"无法读取这个 DOCX：{exc}") from exc


def paragraph_has_toc_field(paragraph) -> bool:
    return bool(paragraph._p.xpath('.//w:instrText[contains(text(), "TOC")]'))


def paragraph_has_pageref_field(paragraph) -> bool:
    return bool(paragraph._p.xpath('.//w:instrText[contains(text(), "PAGEREF")]'))


def find_first_chapter_index(doc: Document) -> Optional[int]:
    return find_first(doc, lambda p: p.style.name.startswith("Heading") and re.match(r"^第\s*(?:1|一)\s*章", p.text.strip()))


def visible_word_count(text: str) -> tuple[int, int, int, int]:
    """Approximate VIP's visible-word statistics using CJK chars, English words and numbers."""
    cjk = len(re.findall(r"[\u3400-\u9fff]", text or ""))
    english_words = len(re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?", text or ""))
    numbers = len(re.findall(r"\d+(?:\.\d+)?", text or ""))
    return cjk + english_words + numbers, cjk, english_words, numbers


def body_text_with_tables(ctx: AuditContext, start: int, end: int) -> str:
    parts = []
    paragraph_index = 0
    for kind, block in ctx.body_blocks:
        if kind == "p":
            if start <= paragraph_index < end:
                parts.append(block.text)
            paragraph_index += 1
        elif start <= paragraph_index < end:
            parts.append(table_plain_text(block))
    return "\n".join(parts)


def normalize_title_for_compare(text: str) -> str:
    text = re.sub(r"[\t\u3000]+", " ", text or "")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s*\d+\s*$", "", text).strip()
    return clean_text(text)


def looks_like_foreign_reference(text: str) -> bool:
    sample = re.sub(r"\[[A-Z]+(?:/OL)?\]", "", text or "")
    letters = len(re.findall(r"[A-Za-z]", sample))
    cjk = len(re.findall(r"[\u3400-\u9fff]", sample))
    return letters >= 20 and letters > cjk * 2


def audit_basic(ctx: AuditContext):
    doc = ctx.doc
    if not doc.paragraphs:
        ctx.add("文件结构", "FAIL", "全文", "文档没有可读取段落。", severity="严重")
        return
    ctx.add(
        "文件结构",
        "PASS",
        "全文",
        "DOCX 可读取。",
        f"段落 {len(doc.paragraphs)}；表格 {len(doc.tables)}；图片 {len(doc.inline_shapes)}；节 {len(doc.sections)}；媒体文件 {ctx.media_count}。",
    )


def audit_template_residue(ctx: AuditContext):
    needles = ["C++语言", "面向对象", "数据结构的重要算法", "…….", "存在着许多不足之处", "线性表的基本理论知识", "00P"]
    hits = []
    for i, p in enumerate(ctx.doc.paragraphs):
        if any(n in p.text for n in needles):
            hits.append((i + 1, p.text.strip()[:160]))
    if hits:
        evidence = "<br>".join(f"P{idx}: {html.escape(text)}" for idx, text in hits[:20])
        ctx.add("模板残留", "FAIL", "全文", "发现疑似模板示例文字残留。", evidence, severity="严重")
    else:
        ctx.add("模板残留", "PASS", "全文", "未发现常见模板示例文字残留。")


def audit_abstracts(ctx: AuditContext):
    doc = ctx.doc
    cn_title = ctx.find_exact("摘  要")
    en_title = ctx.find_exact("Abstract")
    cn_kw = find_first(doc, lambda p: p.text.strip().startswith("关键词"))
    en_kw = find_first(doc, lambda p: p.text.strip().startswith("Keywords"))
    vip_abstract_font_bad = []
    vip_abstract_size_bad = []

    def collect_abstract_font_size(line_no: int, part: str, issues: list[str], sample: str):
        for issue in issues:
            if issue.startswith("字号"):
                vip_abstract_size_bad.append((line_no, part, issue, "应为小四 12 pt", sample[:90]))
            elif "字体" in issue:
                target = "Times New Roman" if issue.startswith("西文字体") else "宋体"
                if part.startswith("英文"):
                    target = "Times New Roman"
                vip_abstract_font_bad.append((line_no, part, issue, target, sample[:90]))

    if cn_title is not None and cn_title > 0:
        title_idx = None
        for idx in range(cn_title - 1, max(-1, cn_title - 8), -1):
            if doc.paragraphs[idx].text.strip():
                title_idx = idx
                break
        if title_idx is not None:
            p = doc.paragraphs[title_idx]
            size = effective_run_size(p)
            east = effective_run_east_asia(p)
            line = effective_line_spacing(p)
            centered = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER
            line_ok = line in (1.5, None)
            before_ok = effective_half_line_spacing(p, "beforeLines")
            after_ok = effective_half_line_spacing(p, "afterLines")
            ok = size == 18 and east == "黑体" and centered and line_ok and before_ok and after_ok
            ctx.add(
                "摘要页论文题目",
                "PASS" if ok else "FAIL",
                f"P{title_idx+1}",
                "摘要页论文题目应为小二号黑体居中，1.5倍行距，段前、段后0.5行间距。",
                f"text={html.escape(p.text.strip())}, size={size}, eastAsia={east}, centered={centered}, lineSpacing={line}, beforeLines0.5={before_ok}, afterLines0.5={after_ok}",
                "重要",
            )
        else:
            ctx.add("摘要页论文题目", "FAIL", "中文摘要", "未在“摘  要”前定位到摘要页论文题目。", severity="重要")

    if cn_title is None:
        ctx.add("中文摘要标题", "FAIL", "中文摘要", "未找到“摘  要”标题。", severity="严重")
    else:
        p = doc.paragraphs[cn_title]
        size = effective_run_size(p)
        east = effective_run_east_asia(p)
        line = effective_line_spacing(p)
        centered = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER
        line_ok = line in (1.5, None)
        before_ok = effective_half_line_spacing(p, "beforeLines")
        after_ok = effective_half_line_spacing(p, "afterLines")
        ok = size == 16 and east == "黑体" and centered and line_ok and before_ok and after_ok
        ctx.add(
            "中文摘要标题",
            "PASS" if ok else "FAIL",
            f"P{cn_title+1}",
            "“摘  要”应为三号黑体居中，中间空两个空格，1.5倍行距，段前、段后0.5行。",
            f"size={size}, eastAsia={east}, centered={centered}, lineSpacing={line}, beforeLines0.5={before_ok}, afterLines0.5={after_ok}",
            "重要",
        )

    if cn_title is not None and cn_kw is not None and cn_kw > cn_title:
        bad_format = []
        for idx in range(cn_title + 1, cn_kw):
            p = doc.paragraphs[idx]
            if not p.text.strip() or has_drawing(p):
                continue
            line_ok = effective_line_spacing(p) in (1.5, None)
            indent_ok = effective_first_line_chars(p, 200)
            font_issues = paragraph_font_issues(p, expected_size=12, expected_east="宋体", expected_latin="Times New Roman")
            collect_abstract_font_size(idx + 1, "中文摘要正文", font_issues, p.text.strip())
            if not line_ok or not indent_ok or font_issues:
                bad_format.append((idx + 1, line_ok, indent_ok, font_issues[:5], p.text.strip()[:50]))
        ctx.add(
            "中文摘要正文格式",
            "PASS" if not bad_format else "FAIL",
            "中文摘要",
            "中文摘要正文应为小四号宋体，1.5倍行距，每段首行缩进2字符；字母、数字使用Times New Roman。",
            "异常项：" + html.escape(str(bad_format[:10])) if bad_format else "未发现摘要正文格式异常。",
            "重要",
        )
        cn = "".join(p.text.strip() for p in doc.paragraphs[cn_title + 1 : cn_kw] if p.text.strip())
        cn_len = len(clean_text(cn))
        if cn_len <= 300:
            ctx.add("中文摘要长度", "PASS", "中文摘要", "中文摘要长度不超过 300 字。", f"当前约 {cn_len} 字。")
        else:
            ctx.add("中文摘要长度", "WARN", "中文摘要", "中文摘要一般不超过 300 字，当前偏长。", f"当前约 {cn_len} 字。", "一般")
    else:
        ctx.add("中文摘要长度", "WARN", "中文摘要", "无法定位中文摘要正文或关键词行。")

    if cn_kw is not None:
        p = doc.paragraphs[cn_kw]
        text = p.text.strip()
        kws = text.split("：", 1)[1] if "：" in text else text.split(":", 1)[1] if ":" in text else ""
        count = len([x for x in re.split(r"[；;]", kws) if x.strip()])
        uses_cn_colon = "关键词：" in text
        uses_cn_semicolon = "；" in kws and ";" not in kws
        trailing = text.endswith(("。", ".", "；", ";"))
        first_run = first_text_run(p)
        keyword_bold = first_run is not None and first_run.text.startswith("关键词") and first_run.font.bold is True
        line = effective_line_spacing(p)
        line_ok = line in (1.5, None)
        indent_ok = effective_first_line_chars(p, 200)
        justify_ok = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.JUSTIFY
        blank_before = cn_kw > 0 and not doc.paragraphs[cn_kw - 1].text.strip()
        font_issues = paragraph_font_issues(p, expected_size=12, expected_east="宋体", expected_latin="Times New Roman")
        collect_abstract_font_size(cn_kw + 1, "中文关键词", font_issues, text)
        ok = (
            3 <= count <= 5
            and not trailing
            and uses_cn_colon
            and uses_cn_semicolon
            and keyword_bold
            and line_ok
            and indent_ok
            and justify_ok
            and blank_before
            and not font_issues
        )
        ctx.add(
            "中文关键词",
            "PASS" if ok else "FAIL",
            f"P{cn_kw+1}",
            "摘要正文与关键词之间应空一行；关键词行应首行缩进2字符、小四号宋体，“关键词”加粗，与内容之间用中文冒号，关键词之间用中文分号，末尾无标点，1.5倍行距，两端对齐，关键词3-5个。",
            f"关键词数={count}, 中文冒号={uses_cn_colon}, 中文分号={uses_cn_semicolon}, 末尾标点={trailing}, 加粗={keyword_bold}, lineSpacing={line}, firstLine2={indent_ok}, justify={justify_ok}, blankBefore={blank_before}, fontIssues={font_issues[:5]}",
            "重要",
        )

    if en_title is None:
        ctx.add("英文摘要标题", "FAIL", "英文摘要", "未找到 Abstract 标题。", severity="严重")
    else:
        title_idx = None
        for idx in range(en_title - 1, max(-1, en_title - 8), -1):
            if doc.paragraphs[idx].text.strip():
                title_idx = idx
                break
        if title_idx is not None:
            tp = doc.paragraphs[title_idx]
            title_size = effective_run_size(tp)
            title_latin = effective_run_latin_font(tp)
            title_bold = effective_run_bold(tp)
            title_line = effective_line_spacing(tp)
            title_centered = effective_paragraph_alignment(tp) == WD_ALIGN_PARAGRAPH.CENTER
            title_line_ok = title_line in (1.5, None)
            title_before_ok = effective_half_line_spacing(tp, "beforeLines")
            title_after_ok = effective_half_line_spacing(tp, "afterLines")
            title_ok = (
                title_size == 18
                and title_latin == "Times New Roman"
                and title_bold is True
                and title_centered
                and title_line_ok
                and title_before_ok
                and title_after_ok
            )
            ctx.add(
                "英文摘要论文题目",
                "PASS" if title_ok else "FAIL",
                f"P{title_idx+1}",
                "Abstract上方英文题目应为Times New Roman、小二、加粗、居中，1.5倍行距，段前段后0.5行。",
                f"text={html.escape(tp.text.strip())}, size={title_size}, latinFont={title_latin}, bold={title_bold}, centered={title_centered}, lineSpacing={title_line}, beforeLines0.5={title_before_ok}, afterLines0.5={title_after_ok}",
                "重要",
            )
        else:
            ctx.add("英文摘要论文题目", "FAIL", "英文摘要", "未在 Abstract 上方定位到英文论文题目。", severity="重要")

        p = doc.paragraphs[en_title]
        size = effective_run_size(p)
        latin = effective_run_latin_font(p)
        bold = effective_run_bold(p)
        line = effective_line_spacing(p)
        centered = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER
        line_ok = line in (1.5, None)
        before_ok = effective_half_line_spacing(p, "beforeLines")
        after_ok = effective_half_line_spacing(p, "afterLines")
        ok = size == 16 and latin == "Times New Roman" and bold is True and centered and line_ok and before_ok and after_ok
        ctx.add(
            "英文摘要标题",
            "PASS" if ok else "FAIL",
            f"P{en_title+1}",
            "Abstract 应为 Times New Roman 三号加粗居中，1.5倍行距，段前、段后0.5行。",
            f"size={size}, latinFont={latin}, bold={bold}, centered={centered}, lineSpacing={line}, beforeLines0.5={before_ok}, afterLines0.5={after_ok}",
            "重要",
        )

    if en_title is not None and en_kw is not None and en_kw > en_title:
        bad_format = []
        for idx in range(en_title + 1, en_kw):
            p = doc.paragraphs[idx]
            if not p.text.strip() or has_drawing(p):
                continue
            line_ok = effective_line_spacing(p) in (1.5, None)
            indent_ok = effective_first_line_chars(p, 200)
            before_ok = effective_half_line_spacing(p, "beforeLines")
            after_ok = effective_half_line_spacing(p, "afterLines")
            font_issues = paragraph_font_issues(p, expected_size=12, expected_east="Times New Roman", expected_latin="Times New Roman")
            collect_abstract_font_size(idx + 1, "英文摘要正文", font_issues, p.text.strip())
            if not line_ok or not indent_ok or not before_ok or not after_ok or font_issues:
                bad_format.append((idx + 1, line_ok, indent_ok, before_ok, after_ok, font_issues[:5], p.text.strip()[:50]))
        ctx.add(
            "英文摘要正文格式",
            "PASS" if not bad_format else "FAIL",
            "英文摘要",
            "英文摘要正文应为小四号Times New Roman，1.5倍行距，每段首行缩进2字符，段前段后0.5行。",
            "异常项：" + html.escape(str(bad_format[:10])) if bad_format else "未发现英文摘要正文格式异常。",
            "重要",
        )
        en = " ".join(p.text.strip() for p in doc.paragraphs[en_title + 1 : en_kw] if p.text.strip())
        words = len(re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?|\d+(?:\.\d+)?", en))
        ctx.add("英文摘要长度", "PASS" if words <= 250 else "WARN", "英文摘要", "英文摘要一般不超过 250 个实词。", f"当前约 {words} 个词。")

    if en_kw is not None:
        p = doc.paragraphs[en_kw]
        text = p.text.strip()
        content = text.split(":", 1)[1] if ":" in text else text.split("：", 1)[1] if "：" in text else ""
        count = len([x for x in re.split(r"[;；]", content) if x.strip()])
        uses_cn_colon = text.startswith("Keywords：")
        uses_cn_semicolon = "；" in content and ";" not in content
        keywords_bold = leading_label_bold(p, "Keywords")
        trailing = text.endswith((".", ";", "；", "。"))
        blank_before = en_kw > 0 and not doc.paragraphs[en_kw - 1].text.strip()
        line = effective_line_spacing(p)
        indent_ok = effective_first_line_chars(p, 200)
        line_ok = line in (1.5, None)
        before_ok = effective_half_line_spacing(p, "beforeLines")
        after_ok = effective_half_line_spacing(p, "afterLines")
        font_issues = paragraph_font_issues(p, expected_size=12, expected_east="Times New Roman", expected_latin="Times New Roman")
        collect_abstract_font_size(en_kw + 1, "英文关键词", font_issues, text)
        ok = (
            uses_cn_colon
            and 3 <= count <= 5
            and uses_cn_semicolon
            and keywords_bold
            and not trailing
            and blank_before
            and indent_ok
            and line_ok
            and before_ok
            and after_ok
            and not font_issues
        )
        ctx.add(
            "英文关键词",
            "PASS" if ok else "FAIL",
            f"P{en_kw+1}",
            "Keywords与Abstract正文之间应空一行，首行缩进2字符；Keywords加粗；Keywords后使用中文冒号，词间用中文分号，末尾不加标点，3-5个，Times New Roman小四，1.5倍行距，段前段后0.5行。",
            f"关键词数={count}, 中文冒号={uses_cn_colon}, 中文分号={uses_cn_semicolon}, Keywords加粗={keywords_bold}, 末尾标点={trailing}, blankBefore={blank_before}, firstLine2={indent_ok}, lineSpacing={line}, beforeLines0.5={before_ok}, afterLines0.5={after_ok}, fontIssues={font_issues[:5]}",
            "重要",
        )
        title_writing_bad = []
        separator_bad = []
        if not uses_cn_colon:
            actual = "Keywords:（英文）（1个半角空格）" if text.startswith("Keywords: ") else ("Keywords:" if text.startswith("Keywords:") else text.split(maxsplit=1)[0])
            title_writing_bad.append((en_kw + 1, actual, "Keywords：（中文）", text[:120]))
        if ";" in content or not uses_cn_semicolon:
            separator_bad.append((en_kw + 1, "分隔符有误", "；(中文)", text[:120]))
        ctx.add(
            "英文关键词标题写法错误（官方检测兼容）",
            "PASS" if not title_writing_bad else "FAIL",
            f"P{en_kw+1}",
            "维普会单独提示英文关键词标题写法；Keywords 后应使用中文冒号“：”，不要写成半角冒号加空格。",
            "异常项：" + html.escape(str(title_writing_bad)) if title_writing_bad else "未发现 Keywords 标题写法异常。",
            "重要",
        )
        ctx.add(
            "英文关键词分隔符问题（官方检测兼容）",
            "PASS" if not separator_bad else "FAIL",
            f"P{en_kw+1}",
            "维普会单独提示英文关键词分隔符问题；英文关键词之间也应使用中文分号“；”。",
            "异常项：" + html.escape(str(separator_bad)) if separator_bad else "未发现英文关键词分隔符异常。",
            "重要",
        )
    ctx.add(
        "摘要字体问题（官方检测兼容）",
        "PASS" if not vip_abstract_font_bad else "FAIL",
        "摘要",
        "维普会在摘要片段中单独提示字体问题；中文摘要中文应为宋体，英文摘要、英文和数字应为 Times New Roman。",
        "异常项：" + html.escape(str(vip_abstract_font_bad[:20])) if vip_abstract_font_bad else "未发现维普式摘要字体异常。",
        "重要",
    )
    ctx.add(
        "摘要字号问题（官方检测兼容）",
        "PASS" if not vip_abstract_size_bad else "FAIL",
        "摘要",
        "维普会在摘要片段中单独提示字号问题；摘要正文和关键词应为小四号 12 pt。",
        "异常项：" + html.escape(str(vip_abstract_size_bad[:20])) if vip_abstract_size_bad else "未发现维普式摘要字号异常。",
        "重要",
    )


def audit_toc(ctx: AuditContext):
    doc = ctx.doc
    toc_start = find_first(doc, lambda p: clean_text(p.text) == "目录")
    toc_field = find_first(doc, paragraph_has_toc_field)
    has_toc_xml = xml_has_toc_field(ctx.document_xml)
    body_start = ctx.first_chapter_index

    if toc_start is None and toc_field is None and not has_toc_xml:
        ctx.add("目录", "FAIL", "目录页", "未找到目录标题，也未检测到 Word 自动目录域。", severity="严重")
        return

    if toc_start is not None:
        toc_title = doc.paragraphs[toc_start]
        title_text_ok = toc_title.text.strip() == "目  录"
        title_size = effective_run_size(toc_title)
        title_east = effective_run_east_asia(toc_title)
        title_line = effective_line_spacing(toc_title)
        title_centered = effective_paragraph_alignment(toc_title) == WD_ALIGN_PARAGRAPH.CENTER
        title_line_ok = title_line in (1.5, None)
        title_before_ok = effective_half_line_spacing(toc_title, "beforeLines")
        title_after_ok = effective_half_line_spacing(toc_title, "afterLines")
        title_ok = (
            title_text_ok
            and title_size == 16
            and title_east == "黑体"
            and title_centered
            and title_line_ok
            and title_before_ok
            and title_after_ok
        )
        ctx.add(
            "目录标题",
            "PASS" if title_ok else "FAIL",
            f"P{toc_start+1}",
            "目录页“目  录”应为一级标题，三号黑体居中，中间空两个空格，1.5倍行距，段前、段后0.5行。",
            f"text={html.escape(toc_title.text.strip())}, style={toc_title.style.name}, size={title_size}, eastAsia={title_east}, centered={title_centered}, lineSpacing={title_line}, beforeLines0.5={title_before_ok}, afterLines0.5={title_after_ok}",
            "重要",
        )
    else:
        loc = f"P{toc_field+1}" if toc_field is not None else "目录页"
        ctx.add(
            "目录标题",
            "WARN",
            loc,
            "检测到 Word 自动目录域，但未在普通段落中读取到“目  录”标题；请人工确认目录页标题是否为三号黑体居中、段前段后0.5行。",
            "检测到自动目录域 TOC；普通段落中未读取到目录标题文本。",
            "一般",
        )

    toc_anchor = toc_start if toc_start is not None else (toc_field if toc_field is not None else 0)
    if body_start is None or body_start <= toc_anchor:
        if has_toc_xml:
            ctx.add(
                "目录内容",
                "WARN",
                "目录页",
                "已检测到 Word 自动目录域；由于该目录未作为普通段落暴露，脚本无法精确定位目录结束位置，请人工确认目录页内容。",
                f"底层 XML 检测到 TOC；PAGEREF 域 {ctx.document_xml.count('PAGEREF')} 个。",
                "一般",
            )
        else:
            ctx.add("目录", "WARN", "目录页", "已检测到目录，但无法准确定位目录结束位置。")
        return
    if toc_start is None and toc_field is not None:
        pageref_count = sum(1 for p in doc.paragraphs[toc_field:body_start] if paragraph_has_pageref_field(p))
        ctx.add(
            "目录内容",
            "WARN",
            "目录页",
            "检测到 Word 自动目录域；由于目录条目以域代码保存，脚本无法完全读取可见目录文字，建议在 Word 中右键更新域后人工复核目录条目。",
            f"自动目录域位置=P{toc_field+1}；目录页到正文首页之间检测到 PAGEREF 域 {pageref_count} 个。",
            "一般",
        )
        ctx.add(
            "目录条目排版",
            "WARN",
            "目录页",
            "自动目录条目需人工确认：小四宋体、1.5倍行距，一级不缩进，二级左缩进2字符，三级左缩进4字符。",
            "检测到自动目录域，未按普通段落逐条检查目录条目。",
            "一般",
        )
        return
    if toc_start is None and has_toc_xml:
        ctx.add(
            "目录内容",
            "WARN",
            "目录页",
            "检测到 Word 自动目录域；由于目录条目以域代码保存，脚本无法完全读取可见目录文字，建议在 Word 中右键更新域后人工复核目录条目。",
            f"底层 XML 检测到 TOC；PAGEREF 域 {ctx.document_xml.count('PAGEREF')} 个。",
            "一般",
        )
        ctx.add(
            "目录条目排版",
            "WARN",
            "目录页",
            "自动目录条目需人工确认：小四宋体、1.5倍行距，一级不缩进，二级左缩进2字符，三级左缩进4字符。",
            "检测到自动目录域，未按普通段落逐条检查目录条目。",
            "一般",
        )
        return

    toc_text = "\n".join(ctx.paragraph_texts[toc_start:body_start])
    has_abs = any(x in toc_text for x in ["摘要", "Abstract", "ABSTRACT"])
    has_placeholder = any(x in toc_text for x in ["线性表", "00P"])
    if has_abs or has_placeholder:
        ctx.add("目录内容", "FAIL", "目录页", "目录包含不应出现的摘要项或模板占位内容。", f"包含摘要={has_abs}；包含模板占位={has_placeholder}", "严重")
    else:
        ctx.add("目录内容", "PASS", "目录页", "目录未包含摘要/ABSTRACT，未发现模板占位内容。")

    toc_entries = []
    for i in range(toc_start + 1, body_start):
        raw = doc.paragraphs[i].text.strip()
        if not raw:
            continue
        entry = re.sub(r"\.{2,}|\u2026+", " ", raw)
        entry = re.sub(r"\s*\d+\s*$", "", entry).strip()
        if entry and not re.fullmatch(r"\d+", entry):
            toc_entries.append((i + 1, raw, normalize_title_for_compare(entry)))
    body_titles = []
    tail_titles = {"致谢", "参考文献", "附录"}
    for i, p in enumerate(doc.paragraphs[body_start:], body_start):
        text = p.text.strip()
        if not text:
            continue
        norm = normalize_title_for_compare(text)
        if p.style.name in ("Heading 1", "Heading 2", "Heading 3") or clean_text(text) in tail_titles:
            body_titles.append((i + 1, text, norm))
    body_title_norms = {norm for _, _, norm in body_titles if norm}
    toc_title_norms = {norm for _, _, norm in toc_entries if norm}
    toc_missing = [(line, raw) for line, raw, norm in toc_entries if norm and norm not in body_title_norms]
    body_missing = [
        (line, raw) for line, raw, norm in body_titles
        if norm and norm not in toc_title_norms and not re.match(r"^附\s*录", raw)
    ]
    ctx.add(
        "目录与正文标题一致性",
        "PASS" if not toc_missing and not body_missing else "FAIL",
        "目录页",
        "目录条目应与正文一、二、三级标题以及致谢、参考文献等结构标题一致，更新目录后不应残留旧标题或漏列正文标题。",
        (
            ("目录中找不到对应正文标题：" + html.escape(str(toc_missing[:15])) if toc_missing else "")
            + ("<br>正文标题未在目录中出现：" + html.escape(str(body_missing[:15])) if body_missing else "")
            if (toc_missing or body_missing)
            else f"目录条目 {len(toc_entries)} 个，正文标题 {len(body_titles)} 个，未发现不一致。"
        ),
        "重要",
    )

    bad = []
    for i in range(toc_start + 1, body_start):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            continue
        pf = p.paragraph_format
        style = p.style
        level = 1 if style.name in ("toc 1", "TOC 1") else 2 if style.name in ("toc 2", "TOC 2") else 3 if style.name in ("toc 3", "TOC 3") else None
        if level is None:
            bad.append((i + 1, "样式不是toc 1/2/3", style.name, t[:60]))
            continue
        expected_left = {1: 0, 2: 420, 3: 840}[level]
        ppr = p._p.pPr
        ind = ppr.ind if ppr is not None else None
        left = int(ind.get(qn("left"))) if ind is not None and ind.get(qn("left")) else 0
        line = effective_line_spacing(p)
        size = style_size(style)
        east = style_east_asia(style)
        sample = first_text_run(p)
        sample_size = run_size(sample) if sample is not None else None
        sample_east = run_east_asia(sample) if sample is not None else None
        if line != 1.5:
            bad.append((i + 1, "行距", line, t[:60]))
        if abs(left - expected_left) > 30:
            bad.append((i + 1, "缩进", left, f"应为{expected_left}", t[:60]))
        if (sample_size if sample_size is not None else size) != 12:
            bad.append((i + 1, "字号", sample_size if sample_size is not None else size, t[:60]))
        if (sample_east if sample_east is not None else east) not in ("宋体", None):
            bad.append((i + 1, "字体", sample_east if sample_east is not None else east, t[:60]))
    ctx.add(
        "目录条目排版",
        "PASS" if not bad else "FAIL",
        "目录页",
        "目录中所列标题不包括摘要和ABSTRACT，主要为正文标题（不超过三级）及致谢、参考文献、附录等；小四宋体、1.5倍行距，一级不缩进，二级左缩进2字符，三级左缩进4字符。",
        "异常项：" + html.escape(str(bad[:15])) if bad else "未发现目录条目排版异常。",
        "重要",
    )


def audit_headings(ctx: AuditContext):
    doc = ctx.doc
    specs = [("Heading 1", 16, "黑体", "一级标题"), ("Heading 2", 14, "黑体", "二级标题"), ("Heading 3", 12, "黑体", "三级标题")]
    for style_name, expected_size, expected_east, label in specs:
        try:
            style = doc.styles[style_name]
        except KeyError:
            ctx.add(label, "FAIL", "样式", f"缺少 {style_name} 样式。", severity="严重")
            continue
        s_size = style_size(style)
        s_east = style_east_asia(style)
        paras = [(i + 1, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.style.name == style_name]
        direct_bad = []
        effective_bad = []
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == style_name:
                r = first_text_run(p)
                if r is None:
                    continue
                size = run_size(r)
                east = run_east_asia(r)
                # None means inherited from style; accept it if style is correct.
                if size not in (None, expected_size) or east not in (None, expected_east):
                    direct_bad.append((i + 1, size, east, p.text[:80]))
                eff_size = effective_run_size(p)
                eff_east = effective_run_east_asia(p)
                if eff_size != expected_size or eff_east != expected_east:
                    effective_bad.append((i + 1, p.text.strip(), eff_size, eff_east))
        # Some Word documents inherit title size from a base style instead of storing it
        # on the named Heading style. Judge visible/effective paragraphs first.
        ok = not effective_bad and not direct_bad
        ctx.add(
            label,
            "PASS" if ok else "FAIL",
            "全文标题",
            f"{label}应为 {expected_size} pt {expected_east}。",
            f"样式 size={s_size}, eastAsia={s_east}; 段落数={len(paras)}; 异常标题={html.escape(str(effective_bad[:20]))}; 直接异常={html.escape(str(direct_bad[:10]))}",
            "重要",
        )

    h1_bad = []
    h1_blank_bad = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name != "Heading 1" or not text:
            continue
        is_chapter = re.match(r"^第\s*(?:\d+|[一二三四五六七八九十]+)\s*章", text)
        if is_chapter and not re.match(r"^第\s*(?:\d+|[一二三四五六七八九十]+)\s*章 {2}\S", text):
            h1_bad.append((i + 1, "章号与章名之间不是两个空格", text))
        if text in ("致谢", "致  谢") and text != "致  谢":
            h1_bad.append((i + 1, "致谢中间应空两个空格", text))
        if text in ("附录", "附  录") and text != "附  录":
            h1_bad.append((i + 1, "附录中间应空两个空格", text))
        centered = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER
        line = effective_line_spacing(p)
        line_ok = line in (1.5, None)
        before_ok = effective_half_line_spacing(p, "beforeLines")
        after_ok = effective_half_line_spacing(p, "afterLines")
        if not centered or not line_ok or not before_ok or not after_ok:
            h1_bad.append((i + 1, "一级标题段落格式异常", f"centered={centered}, line={line}, before0.5={before_ok}, after0.5={after_ok}", text))
        if is_chapter:
            blank_count = 0
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                blank_count += 1
                j += 1
            if blank_count != 1:
                h1_blank_bad.append((i + 1, text, f"标题后空{blank_count}行", "规范：标题后空1行"))
    ctx.add(
        "一级章节标题格式",
        "PASS" if not h1_bad else "FAIL",
        "正文一级标题",
        "章节标题应写作“第1章  引言”形式，章号与标题中间空两个空格；一级标题三号黑体居中，1.5倍行距，段前、段后0.5行。",
        "异常项：" + html.escape(str(h1_bad[:15])) if h1_bad else "未发现一级章节标题格式异常。",
        "重要",
    )
    ctx.add(
        "一级标题后空行",
        "PASS" if not h1_blank_bad else "FAIL",
        "正文一级标题",
        "每章一级标题后应空一行，再开始二级标题或正文。",
        "异常项：" + html.escape(str(h1_blank_bad[:15])) if h1_blank_bad else "未发现一级标题后缺少空行。",
        "重要",
    )
    ctx.add(
        "标题空行问题（官方检测兼容）",
        "PASS" if not h1_blank_bad else "FAIL",
        "正文标题",
        "维普会单独提示标题空行问题；一级标题后应空 1 行，再开始二级标题或正文。",
        "异常项：" + html.escape(str([(line, title, actual, "标题后空1行") for line, title, actual, _ in h1_blank_bad[:20]])) if h1_blank_bad else "未发现维普式标题空行异常。",
        "重要",
    )

    first_chapter = next(
        ((i + 1, p.text.strip()) for i, p in enumerate(doc.paragraphs)
         if p.style.name == "Heading 1" and re.match(r"^第\s*(?:\d+|[一二三四五六七八九十]+)\s*章", p.text.strip())),
        None,
    )
    intro_warn = []
    if first_chapter and clean_text(first_chapter[1]) != "第一章引言":
        intro_warn.append((first_chapter[0], first_chapter[1], "学校维普官方批注示例要求写作“第一章  引言”"))
    ctx.add(
        "引言标题写法（官方检测兼容）",
        "PASS" if not intro_warn else "FAIL",
        "正文引言标题",
        "学校官方维普检测对引言标题可能按“第一章  引言”判定；若学校最终以维普结果为准，请按该写法修改。",
        "提醒项：" + html.escape(str(intro_warn)) if intro_warn else "未发现与官方维普引言标题写法不一致的问题。",
        "重要",
    )

    vip_heading_writing_bad = []
    if first_chapter and clean_text(first_chapter[1]) != "第一章引言":
        vip_heading_writing_bad.append((first_chapter[0], first_chapter[1], "第一章（2个半角空格）引言"))
    ctx.add(
        "正文标题写法错误（官方检测兼容）",
        "PASS" if not vip_heading_writing_bad else "FAIL",
        "正文标题",
        "维普会单独提示正文引言标题写法；学校模板兼容写法建议使用“第一章  引言”。",
        "异常项：" + html.escape(str(vip_heading_writing_bad)) if vip_heading_writing_bad else "未发现维普式正文标题写法异常。",
        "重要",
    )

    h2_bad = []
    vip_heading_number_bad = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name != "Heading 2" or not text:
            continue
        if re.match(r"^\d+\.\d+", text) and not re.match(r"^\d+\.\d+ {1}\S", text):
            h2_bad.append((i + 1, "序数后应空一格", text))
            vip_heading_number_bad.append((i + 1, "标题序号后应空一格", text))
        size = effective_run_size(p)
        east = effective_run_east_asia(p)
        line = effective_line_spacing(p)
        if size != 14 or east != "黑体" or line != 1.5:
            h2_bad.append((i + 1, "二级标题样式异常", f"size={size}, eastAsia={east}, line={line}", text))
        direct_bad = []
        for r in p.runs:
            if not r.text:
                continue
            r_size = run_size(r)
            r_east = run_east_asia(r)
            if r_size not in (None, 14):
                direct_bad.append(f"字号{r_size}")
            if r_east not in (None, "黑体"):
                direct_bad.append(f"字体{r_east}")
        if direct_bad:
            h2_bad.append((i + 1, "二级标题直接格式异常", direct_bad[:5], text))
    ctx.add(
        "二级标题格式",
        "PASS" if not h2_bad else "FAIL",
        "正文二级标题",
        "二级标题应为序数后空一格写标题，四号黑体，1.5倍行距。",
        "异常项：" + html.escape(str(h2_bad[:20])) if h2_bad else "未发现二级标题格式异常。",
        "重要",
    )

    h3_bad = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name != "Heading 3" or not text:
            continue
        if re.match(r"^\d+\.\d+\.\d+", text) and not re.match(r"^\d+\.\d+\.\d+ {1}\S", text):
            h3_bad.append((i + 1, "序数后应空一格", text))
            vip_heading_number_bad.append((i + 1, "标题序号后应空一格", text))
        size = effective_run_size(p)
        east = effective_run_east_asia(p)
        line = effective_line_spacing(p)
        if size != 12 or east != "黑体" or line != 1.5:
            h3_bad.append((i + 1, "三级标题样式异常", f"size={size}, eastAsia={east}, line={line}", text))
        direct_bad = []
        for r in p.runs:
            if not r.text:
                continue
            r_size = run_size(r)
            r_east = run_east_asia(r)
            if r_size not in (None, 12):
                direct_bad.append(f"字号{r_size}")
            if r_east not in (None, "黑体"):
                direct_bad.append(f"字体{r_east}")
        if direct_bad:
            h3_bad.append((i + 1, "三级标题直接格式异常", direct_bad[:5], text))
    ctx.add(
        "三级标题格式",
        "PASS" if not h3_bad else "FAIL",
        "正文三级标题",
        "三级标题应为序数后空一格写标题，小四号黑体，1.5倍行距。",
        "异常项：" + html.escape(str(h3_bad[:20])) if h3_bad else "未发现三级标题格式异常。",
        "重要",
    )

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name == "Heading 1" and re.match(r"^第\s*\d+\s*章", text):
            chapter = chapter_number_from_text(text)
            if chapter is not None and not re.match(rf"^第\s*{chapter_number_to_chinese(chapter)}\s*章", text):
                vip_heading_number_bad.append((i + 1, f"章节序号建议写作第{chapter_number_to_chinese(chapter)}章", text))
    ctx.add(
        "正文标题序号写法错误（官方检测兼容）",
        "PASS" if not vip_heading_number_bad else "FAIL",
        "正文标题",
        "维普会单独提示标题序号写法；一级标题建议使用中文章序“第一章”，二、三级标题序号后应空一格。",
        "异常项：" + html.escape(str(vip_heading_number_bad[:30])) if vip_heading_number_bad else "未发现维普式标题序号写法异常。",
        "重要",
    )

    heading_order_bad = []
    current_chapter: Optional[int] = None
    h1_nums = []
    h2_by_chapter: dict[int, list[int]] = defaultdict(list)
    h3_by_parent: dict[tuple[int, int], list[int]] = defaultdict(list)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 1":
            chapter = chapter_number_from_text(text)
            if chapter is not None:
                current_chapter = chapter
                h1_nums.append(chapter)
            continue
        if p.style.name == "Heading 2":
            m = re.match(r"^(\d+)\.(\d+)\b", text)
            if not m:
                heading_order_bad.append((i + 1, "二级标题缺少形如1.1的编号", text))
                continue
            chapter, section = int(m.group(1)), int(m.group(2))
            h2_by_chapter[chapter].append(section)
            if current_chapter is not None and chapter != current_chapter:
                heading_order_bad.append((i + 1, f"二级标题章号{chapter}与当前第{current_chapter}章不一致", text))
        if p.style.name == "Heading 3":
            m = re.match(r"^(\d+)\.(\d+)\.(\d+)\b", text)
            if not m:
                heading_order_bad.append((i + 1, "三级标题缺少形如1.1.1的编号", text))
                continue
            chapter, section, subsection = int(m.group(1)), int(m.group(2)), int(m.group(3))
            h3_by_parent[(chapter, section)].append(subsection)
            if current_chapter is not None and chapter != current_chapter:
                heading_order_bad.append((i + 1, f"三级标题章号{chapter}与当前第{current_chapter}章不一致", text))

    if h1_nums:
        expected = list(range(1, max(h1_nums) + 1))
        if sorted(h1_nums) != expected:
            heading_order_bad.append(("一级标题", h1_nums, f"章号应连续为{expected}"))
    for chapter, sections in sorted(h2_by_chapter.items()):
        expected = list(range(1, max(sections) + 1))
        if sorted(sections) != expected:
            heading_order_bad.append((f"第{chapter}章二级标题", sections, f"序号应连续为{expected}"))
    for (chapter, section), subsections in sorted(h3_by_parent.items()):
        expected = list(range(1, max(subsections) + 1))
        if sorted(subsections) != expected:
            heading_order_bad.append((f"{chapter}.{section}三级标题", subsections, f"序号应连续为{expected}"))

    ctx.add(
        "标题编号连续性",
        "PASS" if not heading_order_bad else "FAIL",
        "正文标题",
        "正文标题编号应与所在章节一致，并按一级、二级、三级标题逐级连续递增。",
        "异常项：" + html.escape(str(heading_order_bad[:30])) if heading_order_bad else "未发现标题编号层级或连续性异常。",
        "重要",
    )


def audit_body_and_numbering(ctx: AuditContext):
    doc = ctx.doc
    refs_start = ctx.find_exact("参考文献")
    body_start = ctx.first_chapter_index or 0
    nums = []
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i >= refs_start:
            continue
        if paragraph_has_numbering(p) and p.text.strip():
            nums.append((i + 1, p.text.strip()[:120]))
    ctx.add("自动编号残留", "PASS" if not nums else "FAIL", "全文", "不应存在导致缩进错乱的自动编号残留。", "" if not nums else "<br>".join(f"P{i}: {html.escape(t)}" for i, t in nums[:20]), "重要")

    cap_pat = re.compile(r"^[图表][A-Z]?\d+-?\d*\s+.+")
    bad = []
    body_bad = []
    vip_body_font_bad = []
    vip_body_size_bad = []
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i >= refs_start:
            continue
        text = p.text.strip()
        if not text or i < body_start or p.style.name.startswith(("Heading", "toc", "目录", "Bibliography")):
            continue
        if cap_pat.match(text) and len(text) < 70 and "。" not in text and "，" not in text:
            continue
        r = first_text_run(p)
        if r is None:
            continue
        size = run_size(r)
        if len(text) > 80 and size is not None and abs(size - 12) > 0.2:
            bad.append((i + 1, size, text[:80]))
    ctx.add("正文字号", "PASS" if not bad else "FAIL", "正文", "正文长段应为小四 12 pt。", "异常项：" + str(bad[:20]) if bad else "未发现长正文段字号异常。", "重要")

    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i >= refs_start:
            continue
        if i < body_start or not is_probable_body_paragraph(p):
            continue
        indent_ok = effective_first_line_indent_ok(p)
        font_issues = paragraph_font_issues(p, expected_size=12, expected_east="宋体", expected_latin="Times New Roman")
        if not indent_ok or font_issues:
            body_bad.append((i + 1, indent_ok, font_issues[:5], p.text.strip()[:70]))
        for run in p.runs:
            raw = run.text.strip()
            if not raw:
                continue
            size = effective_size_for_run(p, run)
            east = effective_east_asia_for_run(p, run)
            latin = effective_latin_for_run(p, run)
            if size is not None and abs(size - 12) > 0.2:
                vip_body_size_bad.append((i + 1, f"字号{size:g}", short_sample(p.text, 90)))
                break
            if has_cjk(raw) and east not in (None, "宋体"):
                vip_body_font_bad.append((i + 1, f"中文字体{east}", short_sample(p.text, 90)))
                break
            if has_latin_or_digit(raw) and latin not in (None, "Times New Roman"):
                vip_body_font_bad.append((i + 1, f"西文字体{latin}", short_sample(p.text, 90)))
                break
    ctx.add(
        "正文段落格式",
        "PASS" if not body_bad else "FAIL",
        "正文",
        "论文正文应为小四号宋体，外文和数字使用 Times New Roman，每段首行缩进2字符。",
        "异常项：" + html.escape(str(body_bad[:20])) if body_bad else "未发现正文段落字体或首行缩进异常。",
        "重要",
    )
    ctx.add(
        "正文文本内容-字体问题（官方检测兼容）",
        "PASS" if not vip_body_font_bad else "FAIL",
        "正文文本内容",
        "维普会单独提示正文文本中的字体问题；正文中文应为宋体，英文、数字应为 Times New Roman。",
        "异常项：" + html.escape(str(vip_body_font_bad[:20])) if vip_body_font_bad else "未发现维普式正文文本字体异常。",
        "重要",
    )
    ctx.add(
        "正文文本内容-字号问题（官方检测兼容）",
        "PASS" if not vip_body_size_bad else "FAIL",
        "正文文本内容",
        "维普会单独提示正文文本中的字号问题；正文文本应为小四号 12 pt。",
        "异常项：" + html.escape(str(vip_body_size_bad[:20])) if vip_body_size_bad else "未发现维普式正文文本字号异常。",
        "重要",
    )

    paragraph_bad = []
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i >= refs_start:
            continue
        if i < body_start or not is_body_or_caption_paragraph(p):
            continue
        text = p.text.strip()
        if re.fullmatch(r"[（(]\d+-\d+[）)]", text):
            continue
        if re.fullmatch(r"续表(?:\d+-\d+|[A-Z]-\d+)", text):
            continue
        if re.match(r"^\([a-z]\)|^（[a-z]）", text, flags=re.IGNORECASE):
            continue
        cap_like = re.match(r"^[图表]\d+[-.]\d+\s+", text) and len(text) < 100
        align = effective_paragraph_alignment(p)
        line = effective_line_spacing(p)
        if cap_like:
            align_ok = align == WD_ALIGN_PARAGRAPH.CENTER
        else:
            align_ok = align in (WD_ALIGN_PARAGRAPH.JUSTIFY, None)
        line_ok = line in (1.5, None)
        if not align_ok or not line_ok:
            paragraph_bad.append((i + 1, "图表题" if cap_like else "正文段落", str(align), line, text[:80]))
    ctx.add(
        "段落对齐与行距",
        "PASS" if not paragraph_bad else "FAIL",
        "正文段落",
        "正文段落应两端对齐、1.5倍行距；图题和表题应居中，行距符合模板要求。",
        "异常项：" + html.escape(str(paragraph_bad[:30])) if paragraph_bad else "未发现正文段落对齐或行距异常。",
        "重要",
    )

    subitem_bad = []
    subitem_pat = re.compile(r"^(?:\d+）|（\d+）|[①②③④⑤⑥⑦⑧⑨⑩])")
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i >= refs_start:
            continue
        text = p.text.strip()
        if not text or not subitem_pat.match(text):
            continue
        indent_ok = effective_first_line_chars(p, 200)
        line_ok = effective_line_spacing(p) in (1.5, None)
        font_issues = paragraph_font_issues(p, expected_size=12, expected_east="宋体", expected_latin="Times New Roman")
        bold_runs = [r.text for r in p.runs if r.text and r.font.bold is True]
        if not indent_ok or not line_ok or font_issues or bold_runs:
            subitem_bad.append((i + 1, indent_ok, line_ok, font_issues[:5], bold_runs[:3], text[:80]))
    ctx.add(
        "三级以下分项格式",
        "PASS" if not subitem_bad else "FAIL",
        "正文分项",
        "三级以下标题或分项可用“1）”“（1）”“①”等；内容应为小四号宋体，1.5倍行距，首行缩进2字符。",
        "异常项：" + html.escape(str(subitem_bad[:20])) if subitem_bad else "未发现三级以下分项格式异常。",
        "重要",
    )


def audit_chinese_punctuation(ctx: AuditContext):
    doc = ctx.doc
    refs_start = ctx.find_exact("参考文献")
    body_start = ctx.first_chapter_index or 0
    hits = []
    for i, p in enumerate(doc.paragraphs):
        if i < body_start:
            continue
        if refs_start is not None and i >= refs_start:
            continue
        text = p.text.strip()
        if not text or p.style.name.startswith(("Heading", "toc", "目录", "Bibliography")):
            continue
        if not has_cjk(text):
            continue
        for m in re.finditer(r"(?<=[\u3400-\u9fff]),(?=[\u3400-\u9fff])|(?<=[\u3400-\u9fff]),(?=\s*[\u3400-\u9fff])", text):
            hits.append((i + 1, "半角英文逗号", text[max(0, m.start() - 30): min(len(text), m.end() + 30)]))
            break
        if len(hits) >= 20:
            break
    ctx.add(
        "中文标点符号",
        "PASS" if not hits else "WARN",
        "正文",
        "中文正文中逗号等标点应使用全角中文标点，例如用“，”代替英文半角“,”。",
        "异常片段：" + html.escape(str(hits[:20])) if hits else "未发现中文正文中常见半角逗号问题。",
        "一般",
    )
    ctx.add(
        "标点符号问题（官方检测兼容）",
        "PASS" if not hits else "WARN",
        "正文",
        "维普会单独提示标点符号问题；中文上下文中的半角英文逗号“,”应改为全角中文逗号“，”。",
        "异常项：" + html.escape(str([(line, "使用了半角英文标点符号“,(英文)”", "应使用全角中文标点符号“，(中文)”", sample) for line, _, sample in hits[:20]])) if hits else "未发现维普式标点符号异常。",
        "一般",
    )


def audit_figures_tables(ctx: AuditContext):
    doc = ctx.doc
    official_figure_caption_size = 12
    fig_cap_pat = re.compile(r"^图(\d+)-(\d+)\s+(.+)$")
    dot_fig_cap_pat = re.compile(r"^图(\d+)\.(\d+)\s+(.+)$")
    dot_table_cap_pat = re.compile(r"^表(\d+)\.(\d+)\s+(.+)$")
    cap_pat = re.compile(r"^([图表][A-Z]?\d+-?\d*)\s+(.+)")
    body = ctx.full_text
    body_start = ctx.first_chapter_index or 0
    body_end = ctx.first_tail_index()
    captions = []
    under = []
    bad_size = []
    figure_captions = []
    figure_caption_nums = []
    figure_caption_bad = []
    vip_figure_label_bad = []
    image_caption_bad = []
    figure_number_bad = []
    figure_reference_bad = []
    dot_number_bad = []
    dot_number_seen = set()

    def add_dot_number_bad(item):
        key = (item[0], item[1])
        if key not in dot_number_seen:
            dot_number_seen.add(key)
            dot_number_bad.append(item)

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        dot_m = dot_fig_cap_pat.match(text)
        if dot_m and len(text) < 100:
            add_dot_number_bad((i + 1, f"图{dot_m.group(1)}.{dot_m.group(2)}", f"应写作图{dot_m.group(1)}-{dot_m.group(2)}", text))
        m = fig_cap_pat.match(text)
        if not m:
            continue
        if not is_centered(p) and not (len(text) < 100 and "。" not in text and "，" not in text):
            continue
        if len(text) >= 100 or "。" in text or "，" in text:
            continue
        num = f"图{m.group(1)}-{m.group(2)}"
        figure_captions.append((i + 1, int(m.group(1)), int(m.group(2)), num, text))
        figure_caption_nums.append(num)
        if not is_centered(p):
            figure_caption_bad.append((i + 1, num, "未居中", text))
        if re.search(r"图\d+-\d+\s{2,}", text) or not re.match(r"^图\d+-\d+ [^\s].+", text):
            figure_caption_bad.append((i + 1, num, "图号与图题之间应空一格", text))
        for segment_name, segment_text, segment_runs in caption_run_segments(p, text, num, m.group(3)):
            segment_bad = caption_segment_issues(
                p,
                segment_text,
                segment_runs,
                official_figure_caption_size,
                strict_size=segment_name == "文献标注",
            )
            if segment_bad:
                figure_caption_bad.append((
                    i + 1,
                    num,
                    segment_name,
                    short_sample(segment_text, 60),
                    "；".join(segment_bad),
                    text,
                ))
            if segment_name == "图号标签":
                label_bad = vip_figure_label_issues(p, segment_text, segment_runs)
                if label_bad:
                    vip_figure_label_bad.append((i + 1, num, "序号标签", "；".join(label_bad), text))

    body_no_caption = "\n".join(
        text for p, text in zip(doc.paragraphs, ctx.paragraph_texts)
        if not (fig_cap_pat.match(text.strip()) and is_centered(p) and len(text.strip()) < 100)
    )

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if i < body_start or i >= body_end or not text:
            continue
        for m in re.finditer(r"(?<![A-Za-z0-9])([图表])(\d+)\.(\d+)", text):
            expected = f"{m.group(1)}{m.group(2)}-{m.group(3)}"
            add_dot_number_bad((i + 1, m.group(0), f"应写作{expected}", short_sample(text, 100)))
            break
        if dot_table_cap_pat.match(text) and len(text) < 100:
            tm = dot_table_cap_pat.match(text)
            add_dot_number_bad((i + 1, f"表{tm.group(1)}.{tm.group(2)}", f"应写作表{tm.group(1)}-{tm.group(2)}", text))

    for idx, p in enumerate(doc.paragraphs):
        if not has_drawing(p):
            continue
        text_here = p.text.strip()
        if idx < body_start or idx >= body_end:
            continue
        if text_here:
            continue
        if p.style.name.startswith(("Heading", "toc")):
            continue
        next_idx = None
        for j in range(idx + 1, min(len(doc.paragraphs), idx + 8)):
            if doc.paragraphs[j].text.strip():
                next_idx = j
                break
        next_text = doc.paragraphs[next_idx].text.strip() if next_idx is not None else ""
        next_style = doc.paragraphs[next_idx].style.name if next_idx is not None else ""
        if next_style.startswith(("Heading", "toc")) or re.match(r"^(第\s*\d+\s*章|\d+\.\d+)", next_text):
            continue
        if next_idx is None or not fig_cap_pat.match(next_text) or not is_centered(doc.paragraphs[next_idx]):
            image_caption_bad.append((idx + 1, "图片下方未检测到规范图题", next_idx + 1 if next_idx is not None else None, next_text[:80]))

    by_chapter: dict[int, list[int]] = defaultdict(list)
    for _, ch, seq, _, _ in figure_captions:
        by_chapter[ch].append(seq)
    for ch, seqs in sorted(by_chapter.items()):
        expected = list(range(1, max(seqs) + 1))
        if sorted(seqs) != expected:
            figure_number_bad.append((f"第{ch}章", sorted(seqs), f"应连续为{expected}"))
    dup_nums = [num for num, count in Counter(figure_caption_nums).items() if count > 1]
    if dup_nums:
        figure_number_bad.append(("重复图号", dup_nums, "同一图号只能作为一个图题出现"))

    for _, _, _, num, text in figure_captions:
        if body_no_caption.count(num) < 1:
            figure_reference_bad.append((num, text, "正文未检测到图号引用"))

    ctx.add(
        "正文插图题注",
        "PASS" if not figure_caption_bad and not image_caption_bad and figure_captions else "FAIL",
        "正文插图",
        "所有插图应在图位下方设置图号和图题，图号与图题之间空一格，图题居中；按官方检测结果，图题文字按小四宋体复核。",
        (
            f"检测正文图片 {max(len(doc.inline_shapes) - 1, 0)} 张，规范图题 {len(figure_captions)} 个。"
            + ("<br>图题格式异常：" + html.escape(str(figure_caption_bad[:20])) if figure_caption_bad else "")
            + ("<br>图片缺少下方图题：" + html.escape(str(image_caption_bad[:20])) if image_caption_bad else "")
        ),
        "重要",
    )
    ctx.add(
        "插图编号连续性",
        "PASS" if not figure_number_bad and not dot_number_bad else "FAIL",
        "正文插图",
        "图表编号应按一级标题编排，如第3章第1个图为“图3-1”；正文引用也应使用短横线编号，不能写作“图3.1/表3.1”。",
        (
            ("点号编号异常：" + html.escape(str(dot_number_bad[:20])) if dot_number_bad else "")
            + ("<br>连续性异常：" + html.escape(str(figure_number_bad[:20])) if figure_number_bad else "")
            if (dot_number_bad or figure_number_bad)
            else html.escape(str({ch: sorted(seqs) for ch, seqs in sorted(by_chapter.items())}))
        ),
        "重要",
    )
    figure_sequence_bad = [item for item in dot_number_bad if str(item[1]).startswith("图")]
    ctx.add(
        "图序问题（官方检测兼容）",
        "PASS" if not figure_sequence_bad else "FAIL",
        "正文插图",
        "维普会单独提示图序问题；图号应按章节使用短横线编号，例如“图1-1”，不要写作“图1.1”。",
        "异常项：" + html.escape(str(figure_sequence_bad[:20])) if figure_sequence_bad else "未发现维普式图序异常。",
        "重要",
    )
    ctx.add(
        "图题序号标签格式（官方检测兼容）",
        "PASS" if not vip_figure_label_bad else "FAIL",
        "正文插图",
        "维普会单独检查图题开头“图X-X”序号标签；序号标签应使用宋体、小四号，并与图题整体格式一致。",
        "异常项：" + html.escape(str(vip_figure_label_bad[:20])) if vip_figure_label_bad else "未发现图题序号标签字体或字号异常。",
        "重要",
    )
    ctx.add(
        "插图正文引用",
        "PASS" if not figure_reference_bad else "FAIL",
        "正文插图",
        "文中插图必须在正文中引用，不能只出现图片和图题。",
        "未引用：" + html.escape(str(figure_reference_bad[:20])) if figure_reference_bad else f"已检测到 {len(figure_captions)} 个图号均在正文中出现引用。",
        "重要",
    )

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        m = cap_pat.match(text)
        if m and len(text) < 80 and "。" not in text and "，" not in text:
            num = m.group(1)
            captions.append((i + 1, num, text))
            if len(re.findall(re.escape(num), body)) < 2:
                under.append((i + 1, num, text))
            r = first_text_run(p)
            size = run_size(r) if r else None
            if size is not None and abs(size - 10.5) > 0.2:
                bad_size.append((i + 1, size, text))
    ctx.add("图表题字号", "PASS" if not bad_size else "FAIL", "全文图表", "图题、表题应为五号宋体居中。", "异常项：" + str(bad_size[:20]) if bad_size else f"检测图表题 {len(captions)} 个，未发现字号异常。", "重要")
    ctx.add("图表引用", "PASS" if not under else "FAIL", "全文图表", "正文应引用所有图表。", "未充分引用：" + str(under[:20]) if under else "所有图表编号至少出现两次。", "重要")
    ctx.add("图片数量", "PASS" if len(doc.inline_shapes) > 0 else "WARN", "全文图片", "正文应包含必要图片。", f"inline_shapes={len(doc.inline_shapes)}, media={ctx.media_count}")


def cell_border_val(cell, edge: str) -> Optional[str]:
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    borders = tcPr.find(qn("tcBorders"))
    if borders is None:
        return None
    el = borders.find(qn(edge))
    if el is None:
        return None
    return el.get(qn("val"))


def table_border_val(table, edge: str) -> Optional[str]:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None
    borders = tbl_pr.find(qn("tblBorders"))
    if borders is None:
        return None
    el = borders.find(qn(edge))
    if el is None:
        return None
    return el.get(qn("val"))


def table_style_border_val(table, edge: str) -> Optional[str]:
    style = table.style
    if style is None:
        return None
    for item in iter_style_chain(style):
        element = getattr(item, "_element", None)
        if element is None:
            continue
        tbl_pr = element.find(qn("tblPr"))
        if tbl_pr is None:
            continue
        borders = tbl_pr.find(qn("tblBorders"))
        if borders is None:
            continue
        el = borders.find(qn(edge))
        if el is not None:
            return el.get(qn("val"))
    return None


def border_is_visible(value: Optional[str]) -> bool:
    return value not in (None, "nil", "none")


def audit_tables(ctx: AuditContext):
    doc = ctx.doc
    blocks = ctx.body_blocks
    body_cap_pat = re.compile(r"^表(\d+)-(\d+) [^\s].+")
    dot_body_cap_pat = re.compile(r"^表(\d+)\.(\d+)\s+(.+)")
    app_cap_pat = re.compile(r"^表([A-Z])-(\d+) [^\s].+")
    cont_cap_pat = re.compile(r"^续表((?:\d+-\d+)|(?:[A-Z]-\d+))$")
    body_text_no_captions = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if (body_cap_pat.match(text) or app_cap_pat.match(text)) and is_centered(p) and len(text) < 100:
            continue
        body_text_no_captions.append(p.text)
    body_text_no_captions = "\n".join(body_text_no_captions)

    table_infos = []
    caption_bad = []
    ref_bad = []
    number_bad = []
    three_line_bad = []
    table_text_bad = []
    dot_table_bad = []
    body_by_chapter: dict[int, list[int]] = defaultdict(list)
    app_by_letter: dict[str, list[int]] = defaultdict(list)
    caption_nums = []

    for bi, (kind, block) in enumerate(blocks):
        if kind != "t":
            continue
        table = block
        if is_formula_table(table) or is_cover_table(table):
            continue
        prev_p = None
        for j in range(bi - 1, -1, -1):
            if blocks[j][0] == "p" and blocks[j][1].text.strip():
                prev_p = blocks[j][1]
                break
        cap_text = prev_p.text.strip() if prev_p is not None else ""
        dot_m = dot_body_cap_pat.match(cap_text)
        if dot_m:
            dot_table_bad.append((bi + 1, f"表{dot_m.group(1)}.{dot_m.group(2)}", f"应写作表{dot_m.group(1)}-{dot_m.group(2)}", cap_text))
        body_m = body_cap_pat.match(cap_text)
        app_m = app_cap_pat.match(cap_text)
        cont_m = cont_cap_pat.match(cap_text)
        if cont_m:
            num = f"表{cont_m.group(1)}"
            table_infos.append((bi + 1, num, cap_text, len(table.rows), len(table.columns)))
            if not is_centered(prev_p) and effective_paragraph_alignment(prev_p) != WD_ALIGN_PARAGRAPH.RIGHT:
                caption_bad.append((num, "续表标注应位于表格上方右侧", cap_text))
            continue
        if not body_m and not app_m:
            caption_bad.append((bi + 1, "表格上方未检测到规范表题", cap_text[:80]))
            continue

        if body_m:
            num = f"表{body_m.group(1)}-{body_m.group(2)}"
            body_by_chapter[int(body_m.group(1))].append(int(body_m.group(2)))
        else:
            num = f"表{app_m.group(1)}-{app_m.group(2)}"
            app_by_letter[app_m.group(1)].append(int(app_m.group(2)))
        caption_nums.append(num)
        table_infos.append((bi + 1, num, cap_text, len(table.rows), len(table.columns)))

        if not is_centered(prev_p):
            caption_bad.append((num, "表题未居中", cap_text))
        if re.search(r"表(?:\d+-\d+|[A-Z]-\d+)\s{2,}", cap_text):
            caption_bad.append((num, "表号与表题之间应空一格", cap_text))
        r = first_text_run(prev_p)
        size = run_size(r) if r else None
        east = run_east_asia(r) if r else None
        if size is not None and abs(size - 10.5) > 0.2:
            caption_bad.append((num, f"表题字号{size}", cap_text))
        if east not in (None, "宋体"):
            caption_bad.append((num, f"表题中文字体{east}", cap_text))
        for run in prev_p.runs:
            if not run.text:
                continue
            r_size = run_size(run)
            r_east = run_east_asia(run)
            r_latin = run_latin_font(run)
            if r_size is not None and abs(r_size - 10.5) > 0.2:
                caption_bad.append((num, f"表题片段字号{r_size}", run.text))
                break
            if has_cjk(run.text) and r_east not in (None, "宋体"):
                caption_bad.append((num, f"表题片段中文字体{r_east}", run.text))
                break
            if has_latin_or_digit(run.text) and r_latin not in (None, "Times New Roman"):
                caption_bad.append((num, f"表题片段西文字体{r_latin}", run.text))
                break
        if body_text_no_captions.count(num) < 1:
            ref_bad.append((num, cap_text, "正文未检测到表号引用"))

        if table.rows:
            tbl_top = table_border_val(table, "top")
            tbl_bottom = table_border_val(table, "bottom")
            style_top = table_style_border_val(table, "top")
            style_inside_h = table_style_border_val(table, "insideH")
            style_bottom = table_style_border_val(table, "bottom")
            top_vals = [cell_border_val(cell, "top") for cell in table.rows[0].cells]
            mid_vals = [cell_border_val(cell, "bottom") for cell in table.rows[0].cells]
            bottom_vals = [cell_border_val(cell, "bottom") for cell in table.rows[-1].cells]
            vertical_vals = []
            for row in table.rows:
                for cell in row.cells:
                    vertical_vals.extend([cell_border_val(cell, "left"), cell_border_val(cell, "right")])
            has_top = border_is_visible(tbl_top) or border_is_visible(style_top) or any(border_is_visible(v) for v in top_vals)
            has_mid = border_is_visible(style_inside_h) or any(border_is_visible(v) for v in mid_vals)
            has_bottom = border_is_visible(tbl_bottom) or border_is_visible(style_bottom) or any(border_is_visible(v) for v in bottom_vals)
            # Table Grid styles keep default vertical borders in the style XML even when
            # the table/cells explicitly suppress them. Only flag vertical lines that
            # are present on the table itself or on cells, not style defaults alone.
            has_vertical = any(
                border_is_visible(v)
                for v in [
                    table_border_val(table, "left"),
                    table_border_val(table, "right"),
                    table_border_val(table, "insideV"),
                    *vertical_vals,
                ]
            )
            if not has_top:
                three_line_bad.append((num, "首行上边线缺失"))
            if len(table.rows) > 1 and not has_mid:
                three_line_bad.append((num, "表头下方横线缺失"))
            if not has_bottom:
                three_line_bad.append((num, "末行下边线缺失"))
            if has_vertical:
                three_line_bad.append((num, "检测到竖线，三线表一般不应使用竖线"))

        cell_bad_seen = set()
        for row_idx, row in enumerate(table.rows, 1):
            for col_idx, cell in enumerate(row.cells, 1):
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if not text:
                        continue
                    for run in p.runs:
                        raw = run.text.strip()
                        if not raw:
                            continue
                        issues = []
                        size = run_size(run)
                        east = run_east_asia(run)
                        latin = run_latin_font(run)
                        if size is not None and not any(abs(size - ok) <= 0.2 for ok in (10.5, 11, 12)):
                            issues.append(f"字号{size}")
                        if has_cjk(raw) and east not in (None, "宋体"):
                            issues.append(f"中文字体{east}")
                        if re.search(r"[A-Za-z0-9]", raw) and latin not in (None, "Times New Roman"):
                            issues.append(f"西文字体{latin}")
                        if issues:
                            key = (num, row_idx, col_idx, tuple(issues), raw[:24])
                            if key not in cell_bad_seen:
                                cell_bad_seen.add(key)
                                table_text_bad.append((num, f"R{row_idx}C{col_idx}", "；".join(issues[:3]), raw[:60]))
                        if len(table_text_bad) >= 40:
                            break
                    if len(table_text_bad) >= 40:
                        break
                if len(table_text_bad) >= 40:
                    break
            if len(table_text_bad) >= 40:
                break

    for ch, seqs in sorted(body_by_chapter.items()):
        expected = list(range(1, max(seqs) + 1))
        if sorted(seqs) != expected:
            number_bad.append((f"第{ch}章", sorted(seqs), f"应连续为{expected}"))
    for letter, seqs in sorted(app_by_letter.items()):
        expected = list(range(1, max(seqs) + 1))
        if sorted(seqs) != expected:
            number_bad.append((f"附录{letter}", sorted(seqs), f"应连续为{expected}"))
    duplicates = [num for num, count in Counter(caption_nums).items() if count > 1]
    if duplicates:
        number_bad.append(("重复表号", duplicates, "同一表号只能作为一个表题出现"))

    ctx.add(
        "表格题注位置与格式",
        "PASS" if table_infos and not caption_bad else "FAIL",
        "正文和附录表格",
        "表号和表题应放在表格上方正中位置，表号后空一格书写表题，表题用5号宋体。",
        (
            f"检测普通表格 {len(table_infos)} 个。"
            + ("<br>异常项：" + html.escape(str(caption_bad[:25])) if caption_bad else "")
        ),
        "重要",
    )
    ctx.add(
        "表格编号连续性",
        "PASS" if not number_bad and not dot_table_bad else "FAIL",
        "正文和附录表格",
        "表号应按一级标题编排，如第3章第1个表为“表3-1”，不能写作“表3.1”，并在同章内依次连续编号。",
        (
            ("点号编号异常：" + html.escape(str(dot_table_bad[:20])) if dot_table_bad else "")
            + ("<br>连续性异常：" + html.escape(str(number_bad[:20])) if number_bad else "")
            if (dot_table_bad or number_bad)
            else html.escape(str({ch: sorted(seqs) for ch, seqs in sorted(body_by_chapter.items())}))
        ),
        "重要",
    )
    ctx.add(
        "表格正文引用",
        "PASS" if not ref_bad else "FAIL",
        "正文和附录表格",
        "文中表格必须在正文或附录说明文字中引用，不能只出现表格和表题。",
        "未引用：" + html.escape(str(ref_bad[:25])) if ref_bad else f"已检测到 {len(table_infos)} 个表号均在正文或附录说明文字中出现引用。",
        "重要",
    )
    ctx.add(
        "三线表结构",
        "PASS" if not three_line_bad else "WARN",
        "正文和附录表格",
        "表格一般采用三线表，必要时可加辅助线；本项检查首线、表头线、底线和明显竖线。",
        "异常项：" + html.escape(str(three_line_bad[:25])) if three_line_bad else "未发现明显非三线表边框结构。",
        "重要",
    )
    ctx.add(
        "三线表要求（官方检测兼容）",
        "PASS" if not three_line_bad else "WARN",
        "正文表格",
        "维普会单独提示三线表要求；普通表格应保留顶线、表头线和底线，删除不必要的竖线或全框线。",
        "异常项：" + html.escape(str(three_line_bad[:25])) if three_line_bad else "未发现维普式三线表异常。",
        "重要",
    )
    ctx.add(
        "表格内容字体字号",
        "PASS" if not table_text_bad else "FAIL",
        "正文和附录表格",
        "表格内文字应与论文模板保持一致，中文一般使用宋体，英文和数字使用 Times New Roman，字号不应明显偏离五号/小四。",
        "异常项：" + html.escape(str(table_text_bad[:30])) if table_text_bad else "未发现表格内容中明确的字体或字号异常。",
        "重要",
    )

    cont_labels = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.fullmatch(r"续表(?:\d+-\d+|[A-Z]-\d+)", text):
            cont_labels.append(text)
    ctx.add(
        "续表标注",
        "PASS" if cont_labels else "WARN",
        "跨页表格",
        "表格允许下页接写，续页表题可省略，但应在右上方写“续表××”，并重复表头。",
        "检测到续表标注：" + html.escape(str(cont_labels)) if cont_labels else "未检测到续表标注；若全文没有跨页表格可忽略。",
        "一般",
    )


def audit_units(ctx: AuditContext):
    doc = ctx.doc
    full_symbol_pairs = [
        ("长度单位 mm", r"\bmm\b", r"毫米"),
        ("时间单位 ms", r"\bms\b", r"毫秒"),
        ("微米单位", r"(?:μm|µm|\bum\b)", r"微米"),
        ("电压单位 V", r"\bV\b", r"伏特"),
        ("电流单位 A", r"\bA\b", r"安培"),
        ("温度单位 ℃", r"℃", r"摄氏度"),
    ]
    text_units = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            text_units.append((f"P{i+1}", p.text))
    for ti, table in enumerate(doc.tables, 1):
        if is_formula_table(table) or is_cover_table(table):
            continue
        text_units.append((f"T{ti}", table_plain_text(table)))

    all_text = "\n".join(t for _, t in text_units)

    def unit_hit_context(text: str, start: int, end: int, radius: int = 46) -> str:
        left = max(0, start - radius)
        right = min(len(text), end + radius)
        snippet = text[left:right].strip().replace("\n", " / ")
        snippet = re.sub(r"\s+", " ", snippet)
        if left > 0:
            snippet = "..." + snippet
        if right < len(text):
            snippet = snippet + "..."
        return snippet

    mixed = []
    for label, symbol_pat, full_pat in full_symbol_pairs:
        has_symbol = bool(re.search(symbol_pat, all_text, flags=re.IGNORECASE))
        has_full = bool(re.search(full_pat, all_text))
        if has_symbol and has_full:
            locs = []
            for loc, text in text_units:
                matches = []
                for m in re.finditer(symbol_pat, text, flags=re.IGNORECASE):
                    matches.append((m.start(), m.end(), m.group(0), "符号写法"))
                for m in re.finditer(full_pat, text):
                    matches.append((m.start(), m.end(), m.group(0), "中文全称"))
                matches.sort(key=lambda x: x[0])
                for start, end, hit, kind in matches:
                    locs.append((loc, kind, hit, unit_hit_context(text, start, end)))
                    if len(locs) >= 10:
                        break
                if len(locs) >= 10:
                    break
            mixed.append((label, "符号和中文全称混用", locs))

    forbidden = []
    forbidden_terms = ["公斤", "公分", "公厘", "丝米", "秒钟", "公升"]
    for loc, text in text_units:
        hits = [term for term in forbidden_terms if term in text]
        if hits:
            forbidden.append((loc, hits, text.strip().replace("\n", " / ")[:140]))

    table_unit_warn = []
    measurement_hint = re.compile(
        r"坐标|偏差|距离|时间|速度|加速度|尺寸|宽|高|面积|视野|步距|间距|比例|当量|硬度|深|功率|电压|电流|温度|滴数|数量|计数|统计|分辨率"
    )
    for ti, table in enumerate(doc.tables, 1):
        if is_formula_table(table) or is_cover_table(table) or not table.rows:
            continue
        text = table_plain_text(table)
        text_no_ip = re.sub(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", "", text)
        if not measurement_hint.search(text_no_ip):
            continue
        has_number = bool(re.search(r"\d", text))
        has_unit = bool(re.search(r"\b(?:mm|ms|V|A|px|pixel)\b|℃|像素|次|张|个|点|行|列|%", text, flags=re.IGNORECASE))
        if has_number and not has_unit:
            table_unit_warn.append((ti, text.replace("\n", " / ")[:120]))

    status = "PASS" if not mixed and not forbidden else ("FAIL" if forbidden else "WARN")
    evidence = []
    if mixed:
        evidence.append("单位写法混用：" + html.escape(str(mixed[:12])))
    if forbidden:
        evidence.append("疑似非法定或不规范单位：" + html.escape(str(forbidden[:12])))
    if table_unit_warn:
        evidence.append("含数字但未明显标出单位/计数口径的表格，建议人工复核：" + html.escape(str(table_unit_warn[:12])))
    ctx.add(
        "计量单位规范",
        status if not table_unit_warn else ("WARN" if status == "PASS" else status),
        "全文",
        "毕业设计(论文)中量的单位应符合我国法定计量单位，基于SI；同一单位全称、简称或符号用法应保持一致。",
        "<br>".join(evidence) if evidence else "未发现常见非法定单位或明显单位写法混用。",
        "重要",
    )


def audit_formulas(ctx: AuditContext):
    doc = ctx.doc
    formula_infos = []
    bad = []
    math_font_bad = []
    formula_font_bad = []
    formula_size_bad = []
    math_font_seen = set()
    missing_formula_nums = []
    chapter_counts: dict[str, list[int]] = {}
    appendix_nums = []
    for pi, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        has_math = bool(p._p.xpath(".//m:oMath")) or bool(p._p.xpath(".//m:oMathPara"))
        standalone_math = has_math and not text
        if not standalone_math:
            continue
        has_num = bool(re.search(r"\((?:\d+-\d+|A-\d+)\)", text))
        if not has_num:
            snippet = text[:80] if text else "公式对象"
            missing_formula_nums.append((pi, snippet))
        for script_tag, kind in ((M_NS + "sup", "sup"), (M_NS + "sub", "sub")):
            for script in p._p.findall(".//" + script_tag):
                for mrun in script.findall(".//" + M_NS + "r"):
                    texts = mrun.findall(".//" + M_NS + "t")
                    mtext = "".join(t.text or "" for t in texts)
                    if not mtext or not re.search(r"[0-9]", mtext):
                        continue
                    rpr = mrun.find(qn("rPr"))
                    rfonts = rpr.find(qn("rFonts")) if rpr is not None else None
                    latin = rfonts.get(qn("ascii")) or rfonts.get(qn("hAnsi")) if rfonts is not None else None
                    if latin not in (None, "Times New Roman"):
                        shown = math_script_display(mtext, kind)
                        key = (pi, shown, latin)
                        if key not in math_font_seen:
                            math_font_seen.add(key)
                            math_font_bad.append((pi, shown, latin, "公式上标/下标数字按官方检测需使用Times New Roman"))
                            formula_font_bad.append((pi, shown, f"字体{latin}", "Times New Roman"))
                        if len(math_font_bad) >= 30:
                            break
                if len(math_font_bad) >= 30:
                    break

    for ti, table in enumerate(doc.tables, 1):
        if not (len(table.rows) == 1 and len(table.columns) == 2):
            continue
        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]
        has_math_obj = any(p._p.xpath(".//m:oMath") or p._p.xpath(".//m:oMathPara") for p in left.paragraphs)
        num = " ".join(p.text.strip() for p in right.paragraphs if p.text.strip())
        if not has_math_obj and not re.fullmatch(r"\((?:\d+-\d+|A-\d+)\)", num):
            continue
        formula_infos.append((ti, num))
        if not has_math_obj:
            bad.append((ti, num, "左侧未识别到公式对象"))
        if not re.fullmatch(r"\((?:\d+-\d+|A-\d+)\)", num):
            bad.append((ti, num, "公式编号应为(章号-序号)或附录(A-序号)"))

        left_center = all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in left.paragraphs)
        right_ok = True
        right_font_bad = []
        for p in right.paragraphs:
            if p.text.strip() and p.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                right_ok = False
            for r in p.runs:
                if not r.text:
                    continue
                size = run_size(r)
                latin = run_latin_font(r)
                if size != 10.5 or latin != "Times New Roman":
                    right_font_bad.append((size, latin, r.text))
                if size != 10.5:
                    formula_size_bad.append((ti, num, f"字号{size or '未识别'}", "5号 10.5 pt", r.text[:40]))
                if latin != "Times New Roman":
                    formula_font_bad.append((ti, num, f"字体{latin or '未识别'}", "Times New Roman", r.text[:40]))
        if not left_center:
            bad.append((ti, num, "公式未居中"))
        if not right_ok:
            bad.append((ti, num, "编号未右对齐"))
        if right_font_bad:
            bad.append((ti, num, "编号字体字号异常", right_font_bad[:3]))

        m = re.fullmatch(r"\((\d+)-(\d+)\)", num)
        if m:
            chapter_counts.setdefault(m.group(1), []).append(int(m.group(2)))
        am = re.fullmatch(r"\(A-(\d+)\)", num)
        if am:
            appendix_nums.append(int(am.group(1)))

    for ch, nums in chapter_counts.items():
        expected = list(range(1, max(nums) + 1))
        if sorted(nums) != expected:
            bad.append((f"第{ch}章", nums, f"编号应连续为{expected}"))
    if appendix_nums and sorted(appendix_nums) != list(range(1, max(appendix_nums) + 1)):
        bad.append(("附录公式", appendix_nums, "附录公式编号应连续"))

    body_text = ctx.full_text
    for _, num in formula_infos:
        bare = num.strip("()")
        ref_forms = [f"式({bare})", f"式{num}"]
        if not any(x in body_text for x in ref_forms):
            # The formula table itself contains the number, so require a textual reference outside the table.
            bad.append((num, "正文未检测到对应公式引用"))
        if re.match(r"\(\d+-\d+\)", num):
            explain_pat = re.compile(rf"式\({re.escape(bare)}\)中|其中，|式中，")
            if not explain_pat.search(body_text):
                bad.append((num, "未检测到变量说明语句"))

    if missing_formula_nums:
        bad.append(("疑似公式缺少编号或编号分行", missing_formula_nums[:30], "公式编号需与公式同行呈现"))
    if math_font_bad:
        bad.append(("公式/数学对象字体", math_font_bad[:30], "数学对象中的字母、数字或上标数字不符合官方Times New Roman检测要求"))

    ctx.add(
        "正文公式格式",
        "PASS" if not bad else "FAIL",
        "正文公式",
        "正文公式应有编号，编号按章连续编排；公式居中，编号加圆括号右对齐，编号为5号Times New Roman；公式中出现的变量应有说明。",
        f"检测到公式表格 {len(formula_infos)} 个：" + html.escape(str(formula_infos)) + ("<br>异常项：" + html.escape(str(bad[:30])) if bad else ""),
        "重要",
    )
    ctx.add(
        "公式编号缺失（官方检测兼容）",
        "PASS" if not missing_formula_nums else "WARN",
        "正文公式",
        "维普会单独提醒公式编号缺失；公式需标注编号，且编号与公式应在同一行呈现。",
        "异常项：" + html.escape(str([(pi, snippet, "当前公式编号缺失或编号与公式分行排版", "公式需标注编号，且编号与公式应在同一行呈现") for pi, snippet in missing_formula_nums[:30]])) if missing_formula_nums else "未发现维普式公式编号缺失或分行排版问题。",
        "一般" if missing_formula_nums else "重要",
    )
    ctx.add(
        "公式字体问题（官方检测兼容）",
        "PASS" if not formula_font_bad else "FAIL",
        "正文公式",
        "维普会在公式或公式编号附近单独提示字体问题；公式编号及数学对象中的英文、数字应使用 Times New Roman。",
        "异常项：" + html.escape(str(formula_font_bad[:20])) if formula_font_bad else "未发现维普式公式字体异常。",
        "重要",
    )
    ctx.add(
        "公式字号问题（官方检测兼容）",
        "PASS" if not formula_size_bad else "FAIL",
        "正文公式",
        "维普会在公式编号附近单独提示字号问题；公式编号应使用 5 号 10.5 pt。",
        "异常项：" + html.escape(str(formula_size_bad[:20])) if formula_size_bad else "未发现维普式公式字号异常。",
        "重要",
    )


def audit_references(ctx: AuditContext):
    doc = ctx.doc
    refs_start = ctx.find_exact("参考文献")
    if refs_start is None:
        ctx.add("参考文献", "FAIL", "参考文献", "未找到参考文献标题。", severity="严重")
        return
    refs_end = ctx.find_exact("附  录")
    if refs_end is None or refs_end <= refs_start:
        refs_end = len(doc.paragraphs)
    ref_nums = []
    bad_font = []
    bad_gbt = []
    bad_type = []
    standard_type_bad = []
    missing_pub_place = []
    vip_ref_font_bad = []
    vip_ref_size_bad = []
    vip_ref_font_seen = set()
    vip_ref_size_seen = set()
    auto_number_refs = []
    ref_language_counts = Counter()
    num_formats = numbering_formats(ctx.path)
    for i in range(refs_start + 1, refs_end):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text:
            continue
        m = re.match(r"\[(\d+)\]", text)
        if m:
            ref_no = m.group(1)
            ref_nums.append(int(m.group(1)))
            types = re.findall(r"\[([A-Z]+(?:/OL)?)\]", text)
            if not types:
                bad_type.append((i + 1, m.group(1), text[:120]))
            if "GB/T " in text and "[S]" not in text:
                standard_type_bad.append((i + 1, m.group(1), text[:120]))
        else:
            num_id = paragraph_num_id(p)
            fmt = num_formats.get(num_id or "")
            ref_no = str(len(auto_number_refs) + 1) if fmt and "[%1]" in fmt else None
            if fmt and "[%1]" in fmt:
                auto_number_refs.append((i + 1, len(auto_number_refs) + 1, text[:120]))
                ref_nums.append(len(auto_number_refs))
                types = re.findall(r"\[([A-Z]+(?:/OL)?)\]", text)
                if not types:
                    bad_type.append((i + 1, str(len(auto_number_refs)), text[:120]))
                if "GB/T " in text and "[S]" not in text:
                    standard_type_bad.append((i + 1, str(len(auto_number_refs)), text[:120]))
        if ref_no:
            ref_language_counts["foreign" if looks_like_foreign_reference(text) else "chinese"] += 1
        if ref_no and re.search(r"\[(?:D|M)\]", text):
            tail = re.split(r"\[(?:D|M)\]", text, maxsplit=1)[-1]
            if not re.search(r"[\.:：][^,，。]{2,}[:：][^,，。]{2,}[,，]\s*\d{4}", tail):
                missing_pub_place.append((i + 1, ref_no, "保存地或保存单位缺失", text[:140]))
        if re.search(r"Gb/t|gb/t|Gb/T", text):
            bad_gbt.append((i + 1, text[:120]))
        r = first_text_run(p)
        if r:
            size = effective_run_size(p)
            east = effective_run_east_asia(p)
            line_ok = effective_line_spacing(p) in (1.5, None)
            has_cn = has_cjk(text)
            if size != 12 or (has_cn and east not in (None, "宋体")) or not line_ok:
                bad_font.append((i + 1, size, east, line_ok, text[:80]))
        for run in p.runs:
            raw = run.text.strip()
            if not raw:
                continue
            size = effective_size_for_run(p, run)
            east = effective_east_asia_for_run(p, run)
            latin = effective_latin_for_run(p, run)
            if size is not None and abs(size - 12) > 0.2:
                key = (i + 1, f"字号{size:g}", text[:60])
                if key not in vip_ref_size_seen:
                    vip_ref_size_seen.add(key)
                    vip_ref_size_bad.append((i + 1, ref_no or "未编号", f"字号{size:g}", "小四 12 pt", text[:90]))
                break
            if has_cjk(raw) and east not in (None, "宋体"):
                key = (i + 1, f"中文字体{east}", text[:60])
                if key not in vip_ref_font_seen:
                    vip_ref_font_seen.add(key)
                    vip_ref_font_bad.append((i + 1, ref_no or "未编号", f"中文字体{east}", "宋体", text[:90]))
                break
            if has_latin_or_digit(raw) and latin not in (None, "Times New Roman"):
                key = (i + 1, f"西文字体{latin}", text[:60])
                if key not in vip_ref_font_seen:
                    vip_ref_font_seen.add(key)
                    vip_ref_font_bad.append((i + 1, ref_no or "未编号", f"西文字体{latin}", "Times New Roman", text[:90]))
                break
    if not ref_nums:
        ctx.add("参考文献编号", "FAIL", "参考文献", "未识别到参考文献编号。", severity="严重")
    else:
        missing = [n for n in range(1, max(ref_nums) + 1) if n not in set(ref_nums)]
        evidence = f"数量={len(ref_nums)}，缺号={missing[:20]}"
        if auto_number_refs:
            evidence += "；检测到Word自动编号参考文献：" + html.escape(str(auto_number_refs[:20]))
        ctx.add("参考文献编号", "PASS" if not missing else "FAIL", "参考文献", "参考文献编号应连续。", evidence, "重要")
    if ref_nums:
        ref_count_ok = len(ref_nums) >= 10
        ctx.add(
            "参考文献数量要求",
            "PASS" if ref_count_ok else "FAIL",
            "参考文献",
            "参考文献数量应不少于10条；维普统计会区分中文文献和外文文献数量。",
            f"总数={len(ref_nums)}；中文={ref_language_counts['chinese']}条；外文={ref_language_counts['foreign']}条。",
            "严重" if not ref_count_ok else "重要",
        )

    # Citation coverage, ignoring [0,1] style intervals.
    body = ctx.text_until(refs_start)
    cites = []
    for m in re.finditer(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]", body):
        raw = m.group(1)
        # Ignore single 0 or 0,1 intervals.
        if raw.strip().startswith("0"):
            continue
        parts = re.split(r"[,，]", raw)
        for part in parts:
            part = part.strip()
            if "-" in part:
                a, b = [int(x.strip()) for x in part.split("-", 1)]
                if a > 0:
                    cites.extend(range(a, b + 1))
            else:
                n = int(part)
                if n > 0:
                    cites.append(n)
    uncited = sorted(set(ref_nums) - set(cites))
    no_ref = sorted(set(cites) - set(ref_nums))
    ctx.add("参考文献引用", "PASS" if not uncited and not no_ref else "FAIL", "参考文献", "所有参考文献应在正文中引用。", f"未引用={uncited}; 引用无文献={no_ref}", "重要")
    first_order = []
    seen_cites = set()
    for n in cites:
        if n not in seen_cites:
            seen_cites.add(n)
            first_order.append(n)
    order_bad = first_order != ref_nums[: len(first_order)] or len(first_order) != len(ref_nums)
    ctx.add(
        "参考文献排序",
        "PASS" if not order_bad else "FAIL",
        "参考文献",
        "参考文献应按正文中首次出现的先后次序排列；同一文献多次出现只使用同一标号。",
        f"正文首次引用顺序={first_order}; 文献列表顺序={ref_nums}" if order_bad else f"正文首次引用顺序与文献列表顺序一致，共 {len(ref_nums)} 条。",
        "重要",
    )

    citation_bad = []
    cite_pat = re.compile(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]")
    for pi, p in enumerate(doc.paragraphs[:refs_start], 1):
        if p.style.name.startswith(("toc", "Heading")):
            continue
        for r in p.runs:
            if not r.text:
                continue
            for m in cite_pat.finditer(r.text):
                raw = m.group(1).strip()
                if raw.startswith("0"):
                    continue
                size = run_size(r)
                latin = run_latin_font(r)
                superscript = r.font.superscript is True
                if size != 12 or latin != "Times New Roman" or not superscript:
                    citation_bad.append((pi, m.group(0), size, latin, superscript, p.text.strip()[:70]))
    ctx.add(
        "参考文献引用格式",
        "PASS" if not citation_bad else "FAIL",
        "正文引用",
        "正文中的参考文献引用标号应为小四号 Times New Roman 字体，并以上标形式显示。",
        "异常项：" + html.escape(str(citation_bad[:20])) if citation_bad else "未发现正文参考文献引用格式异常。",
        "重要",
    )
    ctx.add("参考文献字体", "PASS" if not bad_font else "FAIL", "参考文献", "参考文献应为小四宋体，英文数字 Times New Roman。", "异常项：" + str(bad_font[:20]) if bad_font else "未发现字体异常。", "重要")
    ctx.add(
        "参考文献字体问题（官方检测兼容）",
        "PASS" if not vip_ref_font_bad else "FAIL",
        "参考文献",
        "维普会在参考文献片段中单独提示字体问题；中文应为宋体，英文和数字应为 Times New Roman。",
        "异常项：" + html.escape(str(vip_ref_font_bad[:20])) if vip_ref_font_bad else "未发现维普式参考文献字体异常。",
        "重要",
    )
    ctx.add(
        "参考文献字号问题（官方检测兼容）",
        "PASS" if not vip_ref_size_bad else "FAIL",
        "参考文献",
        "维普会在参考文献片段中单独提示字号问题；参考文献应为小四号 12 pt。",
        "异常项：" + html.escape(str(vip_ref_size_bad[:20])) if vip_ref_size_bad else "未发现维普式参考文献字号异常。",
        "重要",
    )
    ctx.add("GB/T 大小写", "PASS" if not bad_gbt else "FAIL", "参考文献", "国家标准写法应为 GB/T。", "异常项：" + str(bad_gbt[:20]) if bad_gbt else "未发现 Gb/t 等错误写法。")
    ctx.add(
        "参考文献著录信息",
        "PASS" if not missing_pub_place else "WARN",
        "参考文献",
        "学位论文[D]、专著[M]等参考文献应按GB/T 7714—2015补全保存地、保存单位或出版地、出版社等著录信息。",
        "提醒项：" + html.escape(str(missing_pub_place[:20])) if missing_pub_place else "未发现学位论文或专著明显缺少保存地/出版地信息。",
        "一般",
    )
    ctx.add(
        "保存地或保存单位问题（官方检测兼容）",
        "PASS" if not missing_pub_place else "WARN",
        "参考文献",
        "维普会单独提示学位论文或专著的保存地、保存单位问题；请按 GB/T 7714—2015 补全保存地/保存单位或出版地/出版社。",
        "异常项：" + html.escape(str(missing_pub_place[:20])) if missing_pub_place else "未发现维普式保存地或保存单位问题。",
        "一般",
    )
    ref_type_bad = bad_type + standard_type_bad
    ctx.add(
        "参考文献类型标识",
        "PASS" if not ref_type_bad else "FAIL",
        "参考文献",
        "不同类型文献应按 GB/T 7714—2015 标注文献类型，如期刊[J]、会议[C]、专著[M]、标准[S]、电子文献/联机文献[OL]等。",
        "异常项：" + html.escape(str(ref_type_bad[:20])) if ref_type_bad else "未发现缺少类型标识或标准文献类型误用。",
        "重要",
    )


def audit_document_structure(ctx: AuditContext):
    doc = ctx.doc
    cn_abs = ctx.find_exact("摘  要")
    cn_kw = find_first(doc, lambda p: p.text.strip().startswith("关键词"))
    en_abs = ctx.find_exact("Abstract")
    en_kw = find_first(doc, lambda p: p.text.strip().startswith("Keywords"))

    def first_index(predicate) -> Optional[int]:
        for i, paragraph in enumerate(doc.paragraphs):
            if predicate(i, paragraph):
                return i
        return None

    def previous_nonempty(before: Optional[int], after: Optional[int] = None, latin: bool = False) -> Optional[int]:
        if before is None:
            return None
        start = before - 1
        stop = after if after is not None else -1
        for i in range(start, stop, -1):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue
            if "题  目" in text or text.startswith(("本 科", "20")):
                continue
            if latin and not re.search(r"[A-Za-z]{4,}", text):
                continue
            if not latin and not has_cjk(text):
                continue
            return i
        return None

    toc_idx = first_index(lambda _i, p: clean_text(p.text) == "目录" or paragraph_has_toc_field(p))

    found = [
        ("封面", first_index(lambda i, p: "本科毕业设计" in clean_text(p.text) and i < 10)),
        ("诚信声明", first_index(lambda _i, p: "原创性声明" in p.text)),
        ("版权声明", first_index(lambda _i, p: "版权使用授权书" in p.text)),
        ("中文标题", previous_nonempty(cn_abs)),
        ("中文摘要", cn_abs),
        ("中文关键词", cn_kw),
        ("英文标题", previous_nonempty(en_abs, cn_kw, latin=True)),
        ("英文摘要", en_abs),
        ("英文关键词", en_kw),
        ("中文目录", toc_idx),
        ("正文", first_index(lambda _i, p: p.style.name == "Heading 1" and re.match(r"^第\s*(?:\d+|[一二三四五六七八九十]+)\s*章", p.text.strip()))),
        ("致谢", first_index(lambda _i, p: clean_text(p.text) == "致谢")),
        ("参考文献", first_index(lambda _i, p: clean_text(p.text) == "参考文献")),
    ]
    appendix_idx = first_index(lambda _i, p: clean_text(p.text) == "附录")
    if appendix_idx is not None:
        found.append(("附录", appendix_idx))

    missing = [name for name, pos in found if pos is None]
    present = [(name, pos + 1) for name, pos in found if pos is not None]
    order_bad = []
    last_pos = -1
    for name, pos in present:
        if pos < last_pos:
            order_bad.append((name, pos, "位置早于前一结构要素"))
        last_pos = max(last_pos, pos)

    status = "PASS" if not missing and not order_bad else "FAIL"
    ctx.add(
        "结构缺失",
        "PASS" if not missing else "FAIL",
        "全文结构",
        "论文应包含维普模板要求的必备结构要素；附录为非必有结构。",
        "缺失=" + html.escape(str(missing)) if missing else "未发现必备结构缺失。",
        "严重" if missing else "重要",
    )
    ctx.add(
        "结构顺序",
        "PASS" if not missing and not order_bad else "FAIL",
        "全文结构",
        "论文结构顺序应为：封面-诚信声明-版权声明-中文标题-中文摘要-中文关键词-英文标题-英文摘要-英文关键词-中文目录-正文-致谢-参考文献-附录（非必有）。",
        (
            f"当前结构={html.escape(str(present))}"
            + (f"<br>缺失导致顺序不完整={html.escape(str(missing))}" if missing else "")
            + (f"<br>顺序异常={html.escape(str(order_bad))}" if order_bad else "")
        ),
        "严重" if missing or order_bad else "重要",
    )
    ctx.add(
        "结构要素与顺序",
        status,
        "全文结构",
        "论文结构应依次包含封面、诚信声明、版权声明、中文标题、中文摘要、中文关键词、英文标题、英文摘要、英文关键词、中文目录、正文、致谢、参考文献；附录非必有。",
        f"已识别结构={html.escape(str(present))}"
        + (f"<br>缺失={html.escape(str(missing))}" if missing else "")
        + (f"<br>顺序异常={html.escape(str(order_bad))}" if order_bad else ""),
        "严重" if missing or order_bad else "重要",
    )


def audit_word_count_requirements(ctx: AuditContext):
    body_start = ctx.first_chapter_index
    refs_start = ctx.find_exact("参考文献")
    if body_start is None:
        ctx.add("正文字数要求", "WARN", "正文", "未定位到正文第一章，无法统计正文字数。", severity="一般")
        return
    body_end = refs_start if refs_start is not None and refs_start > body_start else ctx.first_tail_index()
    body_text = body_text_with_tables(ctx, body_start, body_end)
    total, cjk, english_words, numbers = visible_word_count(body_text)
    ctx.add(
        "正文字数要求",
        "PASS" if total >= 10000 else "FAIL",
        "正文",
        "按维普结构统计口径，正文应不少于10000字；本脚本近似统计正文段落和表格中的中文字符、英文单词和数字。",
        f"正文约 {total} 字/词/数字；其中中文字符={cjk}，英文词={english_words}，数字={numbers}；统计范围=P{body_start+1} 到 P{body_end}。",
        "严重" if total < 10000 else "重要",
    )


def audit_appendix(ctx: AuditContext):
    doc = ctx.doc
    app = ctx.find_exact("附  录")
    if app is None:
        ctx.add("附录", "WARN", "附录", "未找到“附  录”标题。")
        return
    r = first_text_run(doc.paragraphs[app])
    size = run_size(r) if r else None
    east = run_east_asia(r) if r else None
    app_bad = []
    p_app = doc.paragraphs[app]
    if p_app.style.name != "Heading 1":
        app_bad.append(f"样式为{p_app.style.name}，应为Heading 1")
    if not is_centered(p_app):
        app_bad.append("未居中")
    if size is not None and abs(size - 16) > 0.2:
        app_bad.append(f"字号{size}，应为三号16pt")
    if east not in (None, "黑体"):
        app_bad.append(f"中文字体{east}，应为黑体")
    if p_app.paragraph_format.page_break_before is not True:
        app_bad.append("未设置段前分页，不能稳定保证单独起页")
    ctx.add(
        "附录一级标题",
        "PASS" if not app_bad else "FAIL",
        f"P{app+1}",
        "附录应作为一级标题，中间空两格，并单独起页。",
        "异常项：" + html.escape(str(app_bad)) if app_bad else f"style={p_app.style.name}, size={size}, eastAsia={east}，段前分页已设置。",
        "重要",
    )

    subheads = [p.text.strip() for p in doc.paragraphs[app + 1 :] if p.style.name.startswith("Heading 2") and p.text.strip().startswith("附录")]
    sub_bad = []
    for text in subheads:
        if not re.match(r"^附录 [A-Z]\s+.+", text):
            sub_bad.append(text)
    ctx.add(
        "附录分项",
        "PASS" if subheads and not sub_bad else "FAIL",
        "附录",
        "附录应按大写字母 A、B、C 连续编号，标题写作“附录 A 标题”。",
        "异常项：" + html.escape(str(sub_bad)) if sub_bad else "；".join(subheads[:10]),
        "重要",
    )

    appendix_text = "\n".join(p.text for p in doc.paragraphs[app + 1 :])
    app_refs_bad = []
    if re.search(r"表[A-Z]\d+", appendix_text):
        app_refs_bad.append("附录表号应写作表A-1、表B-1等，不应写作表A1。")
    if re.search(r"图[A-Z]\d+", appendix_text):
        app_refs_bad.append("附录图号应写作图A-1、图B-1等，不应写作图A1。")
    if re.search(r"式\([A-Z]\d+\)|\([A-Z]\d+\)", appendix_text):
        app_refs_bad.append("附录公式编号应写作式(A-1)、(A-1)等，不应写作式(A1)。")
    ctx.add(
        "附录图表公式编号",
        "PASS" if not app_refs_bad else "FAIL",
        "附录",
        "附录中的图、表、公式等应在阿拉伯数字前冠以附录序码，如图A-1、表B-2、式(A-3)。",
        "异常项：" + html.escape(str(app_refs_bad)) if app_refs_bad else "未发现附录图表公式编号缺少短横线的问题。",
        "重要",
    )


def audit_acknowledgement(ctx: AuditContext):
    doc = ctx.doc
    ack_idx = None
    ack = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().replace(" ", "") == "致谢":
            ack_idx = i
            ack = paragraph
            break
    if ack is None:
        ctx.add("致谢页格式", "FAIL", "致谢", "未找到致谢标题。", severity="重要")
        return

    bad = []
    text = ack.text.strip()
    if text != "致  谢":
        bad.append(f"标题应写作“致  谢”，当前为“{text}”")
    if not ack.style.name.startswith("Heading 1"):
        bad.append(f"致谢应使用一级标题样式，当前样式为{ack.style.name}")
    if not is_centered(ack):
        bad.append("致谢标题未居中")
    east = effective_run_east_asia(ack)
    if east not in (None, "黑体"):
        bad.append(f"中文字体应为黑体，当前为{east}")
    size = effective_run_size(ack)
    if size is not None and abs(size - 16) > 0.2:
        bad.append(f"字号应为三号16pt，当前为{size}")
    if ack.paragraph_format.page_break_before is not True:
        bad.append("致谢标题未设置段前分页，不能稳定保证单独起页")

    ctx.add(
        "致谢页格式",
        "PASS" if not bad else "FAIL",
        f"P{ack_idx + 1}",
        "致谢应作为一级标题，中间空两格，并单独起页。",
        "异常项：" + html.escape(str(bad)) if bad else "致谢标题格式和段前分页设置符合要求。",
        "重要",
    )


def audit_headers_footers_structural(ctx: AuditContext):
    doc = ctx.doc
    fixed_headers = {OFFICIAL_FIXED_HEADER, LEGACY_FIXED_HEADER}
    chapter_titles = {
        clean_text(p.text)
        for p in doc.paragraphs
        if p.style.name.startswith("Heading 1") and re.match(r"^第\s*(?:\d+|[一二三四五六七八九十]+)\s*章", p.text.strip())
    }
    header_bad = []
    header_content_bad = []
    header_seen = []

    for si, sec in enumerate(doc.sections, 1):
        header_paras = [p for p in sec.header.paragraphs if p.text.strip()]
        if not header_paras:
            # Cover/statement pages can be without running headers.
            continue

        p = header_paras[0]
        text = p.text.strip()
        norm = clean_text(text)
        is_fixed = text in fixed_headers
        if text == LEGACY_FIXED_HEADER:
            header_bad.append((si, text, f"官方检测要求使用中文括号：{OFFICIAL_FIXED_HEADER}"))
        is_chapter = norm in chapter_titles
        centered = effective_paragraph_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER
        run_bad = []
        for r in p.runs:
            if not r.text:
                continue
            size = run_size(r)
            east = run_east_asia(r) or style_east_asia(p.style)
            if size is not None and abs(size - 10.5) > 0.2:
                run_bad.append(f"字号{size}")
            if not is_kaiti_font(east):
                run_bad.append(f"中文字体{east}")

        header_seen.append(f"S{si}:{html.escape(text)}")
        if not centered:
            header_bad.append((si, text, "未居中"))
        if run_bad:
            header_bad.append((si, text, "；".join(run_bad[:5])))
        if not (is_fixed or is_chapter):
            issue = "页眉内容既不是学校固定页眉，也不是正文对应章名"
            header_bad.append((si, text, issue))
            header_content_bad.append((si, text, issue, f"非正文节应为“{OFFICIAL_FIXED_HEADER}”，正文节应为对应章名"))
        if text == LEGACY_FIXED_HEADER:
            header_content_bad.append((si, text, "页眉内容错误", OFFICIAL_FIXED_HEADER))

    ctx.add(
        "页眉格式与内容",
        "PASS" if not header_bad else "FAIL",
        "页眉",
        f"正文页眉写每一章对应章名；摘要、目录、致谢、参考文献、附录等页眉为“{OFFICIAL_FIXED_HEADER}”，采用楷体_GB2312五号字居中。",
        "已检测：" + "；".join(header_seen) + ("<br>异常项：" + html.escape(str(header_bad[:20])) if header_bad else ""),
        "重要",
    )
    ctx.add(
        "页眉内容（官方检测兼容）",
        "PASS" if not header_content_bad else "FAIL",
        "页眉问题",
        f"维普会单独提示页眉内容错误；摘要、目录、致谢、参考文献、附录等固定页眉应为“{OFFICIAL_FIXED_HEADER}”，正文页眉应为对应章节名。",
        "异常项：" + html.escape(str(header_content_bad[:20])) if header_content_bad else "未发现维普式页眉内容异常。",
        "重要",
    )


def footer_has_page_field(footer) -> bool:
    for paragraph in footer.paragraphs:
        if paragraph._p.xpath('.//w:instrText[contains(text(), "PAGE")]') or paragraph._p.xpath(".//w:fldChar"):
            return True
    return False


def audit_page_numbers(ctx: AuditContext):
    doc = ctx.doc
    bad = []
    alignment_bad = []
    seen = []
    body_section = None
    for idx, sec in enumerate(doc.sections, 1):
        header_text = clean_text("".join(p.text for p in sec.header.paragraphs))
        if re.search(r"第(?:1|一)章", header_text):
            body_section = idx
            break
    if body_section is None:
        body_section = 4
    for si, sec in enumerate(doc.sections, 1):
        footer = sec.footer
        has_page = footer_has_page_field(footer)
        if si < body_section:
            if has_page:
                bad.append((si, "封面/摘要/目录部分不应有页码"))
            continue

        if not has_page:
            bad.append((si, "正文至附录部分应有阿拉伯数字页码"))
            continue

        paragraph = next((p for p in footer.paragraphs if p.text.strip() or p._p.xpath('.//w:instrText[contains(text(), "PAGE")]') or p._p.xpath(".//w:fldChar")), None)
        centered = paragraph is not None and effective_paragraph_alignment(paragraph) == WD_ALIGN_PARAGRAPH.CENTER
        run_bad = []
        if paragraph is not None:
            for r in paragraph.runs:
                size = run_size(r)
                latin = run_latin_font(r)
                if size is not None and abs(size - 10.5) > 0.2:
                    run_bad.append(f"字号{size}")
                if latin is not None and latin != "Times New Roman":
                    run_bad.append(f"字体{latin}")
        pg = sec._sectPr.find(qn("pgNumType"))
        start = pg.get(qn("start")) if pg is not None else None
        if si == body_section and start != "1":
            bad.append((si, f"正文首页页码应从1开始，当前start={start}"))
        if si > body_section and start is not None:
            bad.append((si, f"后续节不应重新设置起始页码，当前start={start}"))
        if not centered:
            bad.append((si, "页码未居中"))
            alignment_bad.append((si, "左对齐", "居中"))
        if run_bad:
            bad.append((si, "；".join(run_bad[:5])))
        seen.append(f"S{si}:PAGE, centered={centered}, start={start}, runIssues={run_bad[:3]}")

    ctx.add(
        "页码格式",
        "PASS" if not bad else "FAIL",
        "页脚",
        "从正文到附录，页脚用阿拉伯数字连续编排页码；页码位于页脚中间，采用Times New Roman五号字体。封面、摘要、目录部分不编写页码。",
        f"正文起始分节=S{body_section}；已检测：" + "；".join(seen) + ("<br>异常项：" + html.escape(str(bad[:20])) if bad else ""),
        "重要",
    )
    ctx.add(
        "页脚对齐方式问题（官方检测兼容）",
        "PASS" if not alignment_bad else "FAIL",
        "页脚问题",
        "维普会单独提示页脚对齐方式问题；正文至附录页码应在页脚居中显示。",
        "异常项：" + html.escape(str(alignment_bad[:20])) if alignment_bad else "未发现维普式页脚对齐异常。",
        "重要",
    )


def audit_vip_compatibility_summary(ctx: AuditContext):
    def collect(keyword: str):
        return [
            f for f in ctx.findings
            if keyword in f.item
            and "官方检测兼容" in f.item
            and f.status != "PASS"
            and f.item != f"{keyword}（官方检测兼容）"
        ]

    for keyword, location, message in (
        ("字体问题", "全文", "维普原始报告会直接以“字体问题”归类；本项汇总摘要、正文、公式、参考文献等位置已识别出的字体异常，便于和维普报告逐项对照。"),
        ("字号问题", "全文", "维普原始报告会直接以“字号问题”归类；本项汇总摘要、正文、公式、参考文献等位置已识别出的字号异常，便于和维普报告逐项对照。"),
        ("对齐方式问题", "页脚问题", "维普页面设置报告会直接以“对齐方式问题”归类；本项汇总页脚页码等位置已识别出的对齐异常。"),
    ):
        hits = collect(keyword)
        status = "FAIL" if any(f.status == "FAIL" for f in hits) else ("WARN" if hits else "PASS")
        evidence = "异常项：" + html.escape(str([(f.item, f.location, evidence_plain_text(f.evidence)[:160]) for f in hits[:20]])) if hits else f"未发现维普式{keyword}。"
        ctx.add(f"{keyword}（官方检测兼容）", status, location, message, evidence, "重要")


def evidence_plain_text(evidence: str) -> str:
    text = html.unescape(evidence or "")
    text = re.sub(r"<br\s*/?>", "；", text)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def unique_matches(pattern: str, text: str, limit: int = 8) -> list[str]:
    seen = []
    for item in re.findall(pattern, text):
        if item not in seen:
            seen.append(item)
        if len(seen) >= limit:
            break
    return seen


def title_matches(text: str, limit: int = 10) -> list[str]:
    patterns = [
        r"第\s*\d+\s*章\s{0,2}[^'\"，,；;\)\]]+",
        r"(?<![\w.])\d+\.\d+\.\d+\s*[\u3400-\u9fffA-Za-z][^'\"，,；;\)\]]+",
        r"(?<![\w.])\d+\.\d+\s*[\u3400-\u9fffA-Za-z][^'\"，,；;\)\]]+",
    ]
    seen = []
    for pattern in patterns:
        for item in re.findall(pattern, text):
            title = re.sub(r"\s+", " ", item).strip()
            if title and title not in seen:
                seen.append(title)
            if len(seen) >= limit:
                return seen
    return seen


def quoted_items(text: str, limit: int = 12) -> list[str]:
    items = []
    for item in re.findall(r"[\"'‘’“”]([^\"'‘’“”]{2,120})[\"'‘’“”]", text):
        clean = re.sub(r"\s+", " ", item).strip()
        if clean and clean not in items:
            items.append(clean)
        if len(items) >= limit:
            break
    return items


def numbered_refs(text: str, limit: int = 20) -> list[str]:
    refs = []
    for item in re.findall(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]", text):
        if item not in refs:
            refs.append(item)
        if len(refs) >= limit:
            break
    return refs


def human_reference_font_points(raw: str) -> list[str]:
    points = []
    entries = re.findall(r"\((\d+),\s*([^,)]*),\s*'([^']*)',\s*(True|False),\s*'([^']*)'\)", raw)
    for para, size, east, line_ok, sample in entries[:12]:
        issues = []
        size = size.strip()
        if size in ("None", ""):
            issues.append("字号未直接识别到，建议选中该条参考文献统一设为小四")
        elif size != "12":
            issues.append(f"当前字号约 {size} pt，应为小四 12 pt")
        if east != "宋体":
            issues.append(f"中文字体当前为 {east or '未识别'}，应为宋体")
        if line_ok != "True":
            issues.append("行距应为 1.5 倍")
        sample = re.sub(r"\s+", " ", sample).strip()
        if len(sample) > 42:
            sample = sample[:42] + "..."
        if issues:
            points.append(f"P{para} 参考文献“{sample}”：{'；'.join(issues)}。")
    return points


def human_body_format_points(raw: str) -> list[str]:
    points = []
    entries = re.findall(r"\((\d+),\s*(True|False),\s*\[([^\]]*)\],\s*'([^']*)'\)", raw)
    for para, indent_ok, issue_raw, sample in entries[:12]:
        issues = []
        if indent_ok != "True":
            issues.append("首行缩进应设为 2 字符")
        size_hits = re.findall(r"字号([\d.]+)", issue_raw)
        for size in size_hits:
            issues.append(f"当前有 {size} pt 文字，应统一为小四 12 pt")
        if "中文字体" in issue_raw:
            issues.append("中文字体应统一为宋体")
        if "西文字体" in issue_raw:
            issues.append("英文和数字应统一为 Times New Roman")
        sample = re.sub(r"\s+", " ", sample).strip()
        if len(sample) > 46:
            sample = sample[:46] + "..."
        if issues:
            points.append(f"P{para} “{sample}”：{'；'.join(issues)}。")
    return points


def human_figure_caption_points(raw: str, limit: int = 12) -> list[str]:
    marker = "图题格式异常："
    if marker not in raw:
        return []
    payload = raw.split(marker, 1)[1]
    for stop in ("；图片缺少下方图题：", "；图片缺少下方图题", "；图片缺少"):
        if stop in payload:
            payload = payload.split(stop, 1)[0]
            break
    try:
        data = ast.literal_eval(payload.strip())
    except Exception:
        return []

    points: list[str] = []
    for item in data:
        if not isinstance(item, tuple):
            continue
        if len(item) >= 6 and item[2] in ("图号标签", "图题文字", "文献标注"):
            para, num, segment_name, segment_text, issue, full_text = item[:6]
            points.append(
                f"P{para} {num} 的{segment_name}“{segment_text}”：{issue}。"
                f"可在 Word 中搜索完整图题：“{full_text}”"
            )
        elif len(item) >= 4:
            para, num, issue, full_text = item[:4]
            points.append(f"P{para} {num}：{issue}。可在 Word 中搜索完整图题：“{full_text}”")
        if len(points) >= limit:
            break
    return points


def human_missing_image_caption_points(raw: str, limit: int = 6) -> list[str]:
    marker = "图片缺少下方图题："
    if marker not in raw:
        return []
    payload = raw.split(marker, 1)[1]
    try:
        data = ast.literal_eval(payload.strip())
    except Exception:
        return []
    points: list[str] = []
    for item in data:
        if not isinstance(item, tuple) or len(item) < 4:
            continue
        image_para, problem, next_para, next_text = item[:4]
        points.append(
            f"P{image_para} 图片下方没有检测到规范图题；下一段 P{next_para} 为“{next_text}”。"
            "图题应写成“图3-1 图题名称”这种形式，并放在图片下方居中。"
        )
        if len(points) >= limit:
            break
    return points


def human_tuple_points(raw: str, label: str, limit: int = 10) -> list[str]:
    points = []
    for m in re.finditer(r"\((P?\d+|S\d+|T\d+|图\d+-\d+|表[A-Z]?\d+-?\d*|第\d+章|附录[A-Z]?)[^)]{0,180}\)", raw):
        chunk = html.unescape(m.group(0))
        clean = re.sub(r"[()\[\]']", "", chunk)
        clean = re.sub(r"\s+", " ", clean).strip()
        if clean and clean not in points:
            points.append(f"{label}：{clean}")
        if len(points) >= limit:
            break
    return points


def human_unit_mixed_points(raw: str, limit: int = 12) -> list[str]:
    marker = "单位写法混用："
    if marker not in raw:
        return []
    payload = raw.split(marker, 1)[1]
    for stop in ("；含数字但未明显标出单位", "；疑似非法定或不规范单位"):
        if stop in payload:
            payload = payload.split(stop, 1)[0]
    try:
        data = ast.literal_eval(payload.strip())
    except Exception:
        return []
    points: list[str] = []
    for unit_label, issue, locs in data:
        for loc_entry in locs:
            if len(loc_entry) >= 4:
                loc, kind, hit, snippet = loc_entry[:4]
                point = f"{unit_label}：{loc} 发现“{hit}”（{kind}），附近文字：“{snippet}”"
            elif len(loc_entry) >= 2:
                loc, snippet = loc_entry[:2]
                point = f"{unit_label}：{loc} 附近文字：“{snippet}”"
            else:
                continue
            if point not in points:
                points.append(point)
            if len(points) >= limit:
                return points
    return points


def paragraph_context_points(raw: str, limit: int = 8) -> list[str]:
    points: list[str] = []
    samples: list[tuple[str, str]] = []
    samples.extend(re.findall(r"\('P(\d+)'\s*,\s*'([^']{8,260})'\)", raw))
    samples.extend(re.findall(r"\bP(\d+)\s*[:：]\s*([^；<]{8,260})", raw))
    for para, sample in samples:
        clean = re.sub(r"\s+", " ", sample).strip()
        clean = clean.strip("，,。；;：:")
        if len(clean) > 54:
            clean = clean[:54]
        point = f"段落P{para}，可在 Word 中搜索片段：“{clean}”"
        if point not in points:
            points.append(point)
        if len(points) >= limit:
            break
    return points


def extract_literal_list_after(raw: str, marker: str):
    if marker not in raw:
        return None
    payload = raw.split(marker, 1)[1].strip()
    start = payload.find("[")
    if start < 0:
        return None
    depth = 0
    quote = ""
    end = None
    for offset, ch in enumerate(payload[start:], start=start):
        if quote:
            if ch == quote and payload[offset - 1] != "\\":
                quote = ""
            continue
        if ch in ("'", '"'):
            quote = ch
            continue
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                end = offset + 1
                break
    if end is None:
        return None
    try:
        return ast.literal_eval(payload[start:end])
    except Exception:
        return None


def clean_display_text(value, limit: int = 80) -> str:
    text = html.unescape(str(value))
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > limit:
        return text[:limit] + "..."
    return text


def loc_label(value) -> str:
    if isinstance(value, int):
        return f"P{value}"
    text = clean_display_text(value, limit=30)
    if text.isdigit():
        return f"P{text}"
    if re.fullmatch(r"S\d+|T\d+|P\d+", text):
        return text
    return text


def bool_issue(flag, issue: str) -> str:
    return "" if flag is True else issue


def format_font_issue_list(issues) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    if not isinstance(issues, list):
        return result
    for issue in issues:
        text = clean_display_text(issue, limit=40)
        size = re.search(r"字号([\d.]+)", text)
        if size:
            formatted = f"当前有 {size.group(1)} pt 文字，应统一为小四 12 pt"
        elif "中文字体" in text:
            formatted = "中文字体应统一为宋体"
        elif "西文字体" in text:
            formatted = "英文和数字应统一为 Times New Roman"
        elif text:
            formatted = text
        else:
            continue
        if formatted not in seen:
            seen.add(formatted)
            result.append(formatted)
    return result


def format_problem_tuple(item, item_name: str = "") -> str:
    if isinstance(item, str):
        return clean_display_text(item, limit=120)
    if not isinstance(item, tuple):
        return clean_display_text(item, limit=120)

    values = list(item)
    if not values:
        return ""

    loc = loc_label(values[0])

    # Header/footer/page-number checks commonly store section number first.
    if item_name == "页眉格式与内容" and len(values) >= 3:
        section = f"S{values[0]}" if isinstance(values[0], int) else loc
        return f"{section} 页眉为“{clean_display_text(values[1], 70)}”：{clean_display_text(values[2], 90)}"
    if item_name == "页码格式" and len(values) >= 2:
        section = f"S{values[0]}" if isinstance(values[0], int) else loc
        issue = clean_display_text(values[1], 100)
        size = re.search(r"字号([\d.]+)", issue)
        if size:
            return f"{section}：页码字号当前为 {size.group(1)} pt，应为五号 10.5 pt。"
        return f"{section}：{issue}"

    # Reference/citation font tuples should be handled before generic title tuples.
    if item_name == "参考文献引用格式" and len(values) >= 6:
        issues = []
        if values[2] != 12:
            issues.append(f"字号当前为 {values[2] or '未识别'}，应为小四 12 pt")
        if values[3] != "Times New Roman":
            issues.append(f"字体当前为 {values[3] or '未识别'}，应为 Times New Roman")
        if values[4] is not True:
            issues.append("引用标号应设置为上标")
        return f"{loc} 引用标号“{clean_display_text(values[1], 20)}”：{'；'.join(issues)}。附近文字：“{clean_display_text(values[5], 54)}”"

    # Abstract/body paragraph tuples: (para, line_ok, indent_ok, font_issues, sample)
    if len(values) == 5 and isinstance(values[1], bool) and isinstance(values[2], bool) and isinstance(values[3], list):
        issues = [
            bool_issue(values[1], "行距应为 1.5 倍"),
            bool_issue(values[2], "首行缩进应为 2 字符"),
            *format_font_issue_list(values[3]),
        ]
        issues = [x for x in issues if x]
        sample = clean_display_text(values[4], 56)
        return f"{loc} “{sample}”：{'；'.join(issues) if issues else '存在段落格式异常'}。"

    # English abstract tuples: (para, line_ok, indent_ok, before_ok, after_ok, font_issues, sample)
    if len(values) == 7 and all(isinstance(v, bool) for v in values[1:5]) and isinstance(values[5], list):
        issues = [
            bool_issue(values[1], "行距应为 1.5 倍"),
            bool_issue(values[2], "首行缩进应为 2 字符"),
            bool_issue(values[3], "段前应为 0.5 行"),
            bool_issue(values[4], "段后应为 0.5 行"),
            *format_font_issue_list(values[5]),
        ]
        issues = [x for x in issues if x]
        sample = clean_display_text(values[6], 56)
        return f"{loc} “{sample}”：{'；'.join(issues) if issues else '存在英文摘要格式异常'}。"

    # Body/subitem tuples with indentation and optional bold/font problems.
    if len(values) >= 4 and isinstance(values[1], bool) and isinstance(values[2], list):
        issues = [bool_issue(values[1], "首行缩进应为 2 字符"), *format_font_issue_list(values[2])]
        issues = [x for x in issues if x]
        sample = clean_display_text(values[3], 56)
        return f"{loc} “{sample}”：{'；'.join(issues) if issues else '存在正文段落格式异常'}。"
    if len(values) >= 6 and isinstance(values[1], bool) and isinstance(values[2], bool) and isinstance(values[3], list):
        issues = [
            bool_issue(values[1], "首行缩进应为 2 字符"),
            bool_issue(values[2], "行距应为 1.5 倍"),
            *format_font_issue_list(values[3]),
        ]
        if values[4]:
            issues.append("分项标题或正文不应额外加粗")
        sample = clean_display_text(values[5], 56)
        return f"{loc} “{sample}”：{'；'.join(x for x in issues if x)}。"

    # Table of contents and title tuples: (para, issue, actual, expected/sample...)
    if len(values) >= 4 and isinstance(values[0], int) and isinstance(values[1], str):
        if str(values[2]).startswith("标题后空"):
            current = clean_display_text(values[2], 30).replace("标题后", "当前标题后")
            return f"{loc} “{clean_display_text(values[1], 58)}”：{current}，应空 1 行。"
        issue = values[1]
        if issue == "缩进":
            sample = clean_display_text(values[4] if len(values) > 4 else "", 54)
            return f"{loc} 目录条目“{sample}”：左缩进当前约 {values[2]}，{values[3]}。"
        if issue == "字号":
            sample = clean_display_text(values[3], 54)
            actual = "未直接识别到" if values[2] is None else f"{values[2]} pt"
            return f"{loc} 目录条目“{sample}”：字号当前为{actual}，应为小四 12 pt。"
        if issue == "行距":
            sample = clean_display_text(values[3], 54)
            return f"{loc} 目录条目“{sample}”：行距当前为 {values[2]}，应为 1.5 倍。"
        if issue == "字体":
            sample = clean_display_text(values[3], 54)
            return f"{loc} 目录条目“{sample}”：中文字体当前为 {values[2] or '未识别'}，应为宋体。"
        title = clean_display_text(values[-1], 58)
        return f"{loc} “{title}”：{clean_display_text(issue, 80)}。"

    if len(values) == 3 and isinstance(values[0], int) and isinstance(values[1], str):
        return f"{loc} “{clean_display_text(values[2], 58)}”：{clean_display_text(values[1], 90)}。"

    # Font-size-only tuples: (para, size, sample)
    if len(values) == 3 and isinstance(values[0], int) and isinstance(values[1], (int, float)):
        return f"{loc} “{clean_display_text(values[2], 58)}”：当前字号约 {values[1]} pt，应为小四 12 pt。"

    # Reference/citation font tuples.
    if item_name == "参考文献字体" and len(values) >= 5:
        issues = []
        if values[1] in (None, ""):
            issues.append("字号未直接识别到，建议统一设为小四 12 pt")
        elif abs(float(values[1]) - 12) > 0.2:
            issues.append(f"当前字号约 {values[1]} pt，应为小四 12 pt")
        if values[2] != "宋体":
            issues.append(f"中文字体当前为 {values[2] or '未识别'}，应为宋体")
        if values[3] is not True:
            issues.append("行距应为 1.5 倍")
        return f"{loc} 参考文献“{clean_display_text(values[4], 58)}”：{'；'.join(issues) if issues else '字体需要人工复核'}。"
    # Figure/table/reference coverage tuples.
    if len(values) >= 3 and re.match(r"^(图|表)[A-Z]?\d+[-.]?\d*", str(values[0])):
        return f"{clean_display_text(values[0], 20)} “{clean_display_text(values[1], 70)}”：{clean_display_text(values[2], 90)}。"
    if len(values) >= 3 and isinstance(values[0], int) and re.match(r"^(图|表)[A-Z]?\d+[-.]?\d*", str(values[1])):
        return f"{loc} {clean_display_text(values[1], 20)} “{clean_display_text(values[2], 70)}”需要在正文中明确引用。"

    # Table/three-line-table/formula tuples where first value is already an item number.
    if len(values) == 2:
        return f"{loc}：{clean_display_text(values[1], 100)}"
    if len(values) >= 3 and not isinstance(values[1], bool):
        return f"{loc}：{clean_display_text(values[1], 70)}；{clean_display_text(values[2], 90)}"

    parts = [str(x) for x in item if x not in (None, "", [])]
    if not parts:
        return ""
    loc = loc_label(parts[0])
    details = "；".join(parts[1:])
    if not details:
        return str(loc)
    return f"{loc}：{details}"


def generic_exception_points(raw: str, item_name: str = "", limit: int = 14) -> list[str]:
    points: list[str] = []
    for marker, label in (
        ("异常项：", "异常项"),
        ("提醒项：", "提醒项"),
        ("需人工复核：", "需复核"),
        ("点号编号异常：", "点号编号异常"),
        ("连续性异常：", "连续性异常"),
        ("图题格式异常：", "图题格式异常"),
        ("图片缺少下方图题：", "图片缺少图题"),
        ("未引用：", "未引用项"),
        ("未充分引用：", "未充分引用项"),
    ):
        data = extract_literal_list_after(raw, marker)
        if not isinstance(data, list):
            continue
        for item in data:
            text = format_problem_tuple(item, item_name)
            if text:
                points.append(f"{label}：{text}")
            if len(points) >= limit:
                return points
    return points


def generic_detected_points(raw: str, limit: int = 8) -> list[str]:
    if "已检测：" not in raw:
        return []
    segment = raw.split("已检测：", 1)[1].split("；异常项：", 1)[0]
    entries = [x.strip() for x in segment.split("；") if x.strip()]
    points = []
    for entry in entries[:limit]:
        points.append(f"已检测到：{format_detected_entry(entry)}")
    return points


def format_detected_entry(entry: str) -> str:
    entry = html.unescape(entry).strip()
    header = re.fullmatch(r"(S\d+):(.+)", entry)
    if header and "PAGE" not in entry:
        return f"{header.group(1)} 页眉为“{clean_display_text(header.group(2), 80)}”"
    page = re.match(r"(S\d+):PAGE,\s*centered=(True|False),\s*start=([^,]+),\s*runIssues=\[([^\]]*)\]", entry)
    if page:
        section, centered, start, run_issues = page.groups()
        bits = [f"{section} 页脚有页码"]
        bits.append("已居中" if centered == "True" else "未居中")
        if start not in ("None", ""):
            bits.append(f"起始页={start}")
        issues = re.findall(r"'([^']+)'", run_issues)
        bits.extend(issues)
        return "，".join(bits)
    return clean_display_text(entry, 120)


def generic_suggestion_points(raw: str, item_name: str) -> list[str]:
    suggestions: list[str] = []
    if "未设置段前分页" in raw or "不能稳定保证单独起页" in raw:
        suggestions.append("在 Word 中把该标题设置为“段前分页”，确保单独起页。")
    if "官方检测要求使用中文括号" in raw:
        suggestions.append(f"把页眉中的英文括号“(论文)”改为中文括号“（论文）”。")
    if "页码未居中" in raw or "centered=False" in raw:
        suggestions.append("把对应节的页脚页码设置为居中。")
    if "字号12.0" in raw and ("页码" in item_name or "页脚" in raw):
        suggestions.append("把页码字号改为五号 10.5pt，英文字体设为 Times New Roman。")
    if "首行上边线缺失" in raw or "表头下方横线缺失" in raw or "末行下边线缺失" in raw:
        suggestions.append("按三线表补齐顶线、表头线和底线，并删除多余竖线。")
    if "保存地或保存单位缺失" in raw:
        suggestions.append("学位论文类参考文献请补全保存地或保存单位信息。")
    if "正文未检测到图号引用" in raw:
        suggestions.append("在正文描述中加入对应图号引用，例如“如图3-1所示”。")
    if "图片下方未检测到规范图题" in raw:
        suggestions.append("在图片下方补充规范图题，例如“图3-1 图题名称”。")
    return suggestions


def human_evidence_points(f: Finding) -> list[str]:
    raw = evidence_plain_text(f.evidence)
    points: list[str] = []

    if f.status == "PASS":
        return ["检查通过，不需要处理。"]

    points.extend(generic_exception_points(raw, f.item))
    points.extend(generic_detected_points(raw))
    points.extend(generic_suggestion_points(raw, f.item))

    if f.item == "参考文献字体":
        points.extend(human_reference_font_points(raw))
        if not points:
            points.append("请选中参考文献列表，统一设置为小四宋体；英文和数字设置为 Times New Roman，行距 1.5 倍。")

    if f.item in ("正文段落格式", "正文字号"):
        points.extend(human_body_format_points(raw))

    if f.item == "计量单位规范":
        points.extend(human_unit_mixed_points(raw))

    titles = title_matches(raw)
    if titles and ("标题" in f.item or "章节" in f.item):
        points.append("需要检查的标题：" + "；".join(titles) + (" 等" if len(titles) >= 10 else ""))

    context_points = [] if f.item == "计量单位规范" and points else paragraph_context_points(raw)
    if context_points:
        points.extend(context_points)

    locs = unique_matches(r"\bP\d+\b", raw)
    if locs and not context_points:
        points.append("涉及位置：" + "、".join(locs) + (" 等" if len(locs) >= 8 else ""))

    figure_nums = unique_matches(r"图\d+-\d+", raw, limit=12)
    if figure_nums and ("图" in f.item or "插图" in f.item):
        points.append("需要检查的图：" + "、".join(figure_nums) + (" 等" if len(figure_nums) >= 12 else ""))

    table_nums = unique_matches(r"表[A-Z]?\d+-\d+|表\d+-\d+", raw, limit=12)
    if table_nums and "表" in f.item:
        points.append("需要检查的表：" + "、".join(table_nums) + (" 等" if len(table_nums) >= 12 else ""))

    refs = numbered_refs(raw)
    if refs and "参考文献" in f.item and f.item != "参考文献字体":
        points.append("涉及参考文献编号：" + "、".join(f'[{r}]' for r in refs) + (" 等" if len(refs) >= 20 else ""))

    if "参考文献" in f.item and ("未识别到参考文献编号" in raw + f.message or "未识别到" in raw + f.message):
        points.append("参考文献列表每条开头应写成 [1]、[2]、[3] 这样的编号；当前没有识别到规范编号。")

    if "图题格式异常" in raw:
        points.extend(human_figure_caption_points(raw))
    if "图片缺少下方图题" in raw:
        points.extend(human_missing_image_caption_points(raw))
    if ("图题格式异常" in raw or "图片缺少下方图题" in raw) and not any("图题" in p or "图号" in p for p in points):
        points.extend(human_tuple_points(raw, "图片/图题问题"))
    if "图题格式异常" in raw or "图片缺少下方图题" in raw:
        points.append("图片下方应有类似“图3-1 图题名称”的图题，图题居中、五号宋体。")

    if "表题" in raw or "表格上方" in raw or "三线表" in f.item:
        points.extend(human_tuple_points(raw, "表格问题"))

    if "页眉" in f.item:
        points.extend(human_tuple_points(raw, "页眉问题"))
    if "页码" in f.item:
        points.extend(human_tuple_points(raw, "页码问题"))

    if "自动编号残留" == f.item:
        snippets = quoted_items(raw, limit=8)
        if snippets:
            points.append("需要检查的文字片段：" + "；".join(snippets) + (" 等" if len(snippets) >= 8 else ""))

    if "未找到" in f.message or "未找到" in raw:
        points.append("文档中没有找到这一项，请先确认是否缺少对应页面或标题。")
    if "未居中" in raw or "centered=False" in raw:
        points.append("需要设置为居中。")
    if "lineSpacing=None" in raw or "lineSpacing=False" in raw or "line=False" in raw:
        points.append("需要把行距设置为 1.5 倍。")
    if "beforeLines0.5=False" in raw:
        points.append("需要设置段前 0.5 行。")
    if "afterLines0.5=False" in raw:
        points.append("需要设置段后 0.5 行。")
    if "firstLine2=False" in raw or "首行缩进" in raw and "False" in raw:
        points.append("需要设置首行缩进 2 字符。")
    if "fontIssues=[]" not in raw and "fontIssues=" in raw:
        points.append("字体存在不符合项，请按要求设置中文字体和英文字体。")
    if "字号" in raw and not any("字号" in p for p in points):
        points.append("字号可能不符合要求，请按本项说明设置。")

    keyword_count = re.search(r"关键词数=(\d+)", raw)
    if keyword_count:
        points.append(f"检测到关键词 {keyword_count.group(1)} 个。")
    if "中文冒号=False" in raw:
        points.append("“关键词”和内容之间应使用中文冒号“：”。")
    if "中文分号=False" in raw:
        points.append("关键词之间应使用中文分号“；”。")
    if "末尾标点=True" in raw:
        points.append("最后一个关键词后面不要加标点。")
    if "加粗=False" in raw or "Keywords加粗=False" in raw:
        points.append("“关键词”或“Keywords”需要加粗。")
    if "blankBefore=False" in raw:
        points.append("正文和关键词之间需要空一行。")

    if "章号与章名之间不是两个空格" in raw:
        points.append("章标题中的章号和标题之间需要空两个空格，例如“第1章  引言”。")
    if "序数后应空一格" in raw:
        points.append("标题序号后面需要空一格，例如“4.7.2 实验结果与误差分析”。")
    if "一级标题后应空一行" in f.message:
        for para, title, actual in re.findall(r"\((\d+),\s*'([^']+)',\s*'标题后空(\d+)行',\s*'规范：标题后空1行'\)", raw):
            points.append(f"P{para} “{title}”：当前标题后空{actual}行，应空1行。")
        points.append("一级标题后面需要先空一行，再开始正文或二级标题。")
    if "未充分引用" in raw or "正文未检测到" in raw:
        points.append("正文中需要明确引用对应的图、表或公式编号。")
    if "未引用=" in raw or "引用无文献=" in raw:
        points.append("参考文献编号和正文引用编号需要一一对应。")
    missing_ref = re.search(r"引用无文献=\[([^\]]*)\]", raw)
    if missing_ref and missing_ref.group(1).strip():
        points.append("正文引用了这些编号，但参考文献列表中没有对应条目：" + missing_ref.group(1).strip())
    uncited_ref = re.search(r"未引用=\[([^\]]*)\]", raw)
    if uncited_ref and uncited_ref.group(1).strip():
        points.append("参考文献列表中这些编号未在正文出现：" + uncited_ref.group(1).strip())
    missing_nums = re.search(r"缺号=\[([^\]]*)\]", raw)
    if missing_nums and missing_nums.group(1).strip():
        points.append("参考文献编号不连续，缺少：" + missing_nums.group(1).strip())
    if f.item == "参考文献排序":
        first = re.search(r"正文首次引用顺序=\[([^\]]*)\]", raw)
        listed = re.search(r"文献列表顺序=\[([^\]]*)\]", raw)
        if first:
            points.append("正文首次出现的引用顺序为：" + first.group(1).strip())
        if listed:
            points.append("参考文献列表当前顺序为：" + (listed.group(1).strip() or "未识别到规范编号"))
        points.append("请按正文第一次引用的先后顺序重新排列参考文献。")
    if "单位写法混用" in raw:
        points.append("同一种单位不要同时使用中文名称和符号写法，请统一。")
        units = unique_matches(r"(?:长度单位 mm|时间单位 ms|微米单位|电压单位 V|电流单位 A|温度单位 ℃)", raw, limit=8)
        if units:
            points.append("涉及单位：" + "、".join(units))
    if "续表" in f.item and "检测到续表标注" in raw:
        points.append("已检测到续表标注。")

    if f.status == "WARN" and not points:
        points.append("这是一项提醒，不一定必须修改；如果学院要求严格，请按说明调整。")
    if f.status == "FAIL" and not points:
        points.append("这一项不符合要求，请按上方说明修改。")

    # Remove duplicates while preserving order.
    result = []
    for point in points:
        if point not in result:
            result.append(point)
    return result


def split_reason_suggestion(f: Finding) -> tuple[list[str], list[str]]:
    points = expand_grouped_points(human_evidence_points(f))
    if f.status == "PASS":
        return points, ["不需要修改。"]

    reason_prefixes = (
        "需要检查",
        "涉及",
        "检测到",
        "文档中没有找到",
        "正文引用了",
        "参考文献列表中",
        "参考文献编号不连续",
        "正文首次出现",
        "参考文献列表当前",
        "图片/图题问题",
        "表格问题",
        "页眉问题",
        "页码问题",
        "异常项",
        "提醒项",
        "需复核",
        "点号编号异常",
        "连续性异常",
        "图题格式异常",
        "图片缺少图题",
        "未引用项",
        "未充分引用项",
        "已检测到",
        "P",
        "S",
        "T",
    )
    suggestion_keywords = (
        "需要设置",
        "应",
        "请",
        "不要",
        "统一",
        "改成",
        "空一格",
        "空两个空格",
        "加粗",
        "居中",
        "勾选",
        "把",
        "补齐",
        "加入",
        "补充",
        "删除",
    )
    reasons: list[str] = []
    suggestions: list[str] = []
    for point in points:
        if point.startswith(reason_prefixes) and not point.startswith(("需要设置", "请按", "请选中")):
            reasons.append(point)
        elif any(key in point for key in suggestion_keywords):
            suggestions.append(point)
        else:
            reasons.append(point)

    if f.status == "WARN" and not suggestions:
        suggestions.append("这是一项提醒，请人工复核；如果学院要求严格，再按本项说明调整。")
    if f.status == "FAIL" and not suggestions:
        suggestions.append("请按本项格式要求修改。")
    if not reasons:
        reasons.append("检测结果显示该项未完全满足格式要求。")

    def dedupe(items: list[str]) -> list[str]:
        result = []
        for item in items:
            if item not in result:
                result.append(item)
        return result

    return dedupe(reasons), dedupe(suggestions)


def expand_grouped_points(points: list[str]) -> list[str]:
    """Split long human-facing checklist points into individually checkable items."""
    split_specs = (
        ("需要检查的标题：", "；"),
        ("涉及位置：", "、"),
        ("需要检查的图：", "、"),
        ("需要检查的表：", "、"),
        ("涉及参考文献编号：", "、"),
        ("需要检查的文字片段：", "；"),
        ("涉及单位：", "、"),
    )
    expanded: list[str] = []
    for point in points:
        matched = False
        for prefix, sep in split_specs:
            if not point.startswith(prefix):
                continue
            body = point[len(prefix):].strip()
            body = re.sub(r"\s*等\s*$", "", body)
            parts = [p.strip() for p in body.split(sep) if p.strip()]
            if len(parts) > 1:
                expanded.extend(prefix + part for part in parts)
                matched = True
            break
        if not matched:
            expanded.append(point)
    result: list[str] = []
    for point in expanded:
        if point not in result:
            result.append(point)
    return result


def checklist_key(*parts: str) -> str:
    data = "\n".join(parts).encode("utf-8", errors="ignore")
    return hashlib.sha1(data).hexdigest()[:16]


def search_snippet_from_point(item: str) -> str:
    m = re.search(r"可在 Word 中搜索(?:片段|完整图题)：“([^”]+)”", item)
    return m.group(1).strip() if m else ""


def render_check_items(card_key: str, group: str, items: list[str], checkable: bool = True) -> str:
    lis = []
    for i, item in enumerate(items, 1):
        key = checklist_key(card_key, group, str(i), item)
        snippet = search_snippet_from_point(item)
        copy_button = ""
        if snippet:
            copy_button = (
                f'<button class="copy-snippet" type="button" '
                f'data-copy-text="{html.escape(snippet, quote=True)}">复制片段</button>'
            )
        if checkable:
            lis.append(
                f"""
                <li class="check-item">
                  <label class="check-row">
                    <input type="checkbox" class="sub-check" data-check-key="{key}">
                    <span class="check-text">{html.escape(item)}</span>
                  </label>
                  {copy_button}
                </li>
                """
            )
        else:
            lis.append(f'<li class="check-item"><span class="check-text">{html.escape(item)}</span>{copy_button}</li>')
    return "<ul class=\"plain-list check-list\">" + "".join(lis) + "</ul>"


def build_report(ctx: AuditContext, out: Path):
    counts = Counter(f.status for f in ctx.findings)
    cards = []
    status_label = {"PASS": "通过", "WARN": "提醒", "FAIL": "未通过"}
    report_key = checklist_key(str(ctx.path))
    embedded_state_json = json.dumps({}, ensure_ascii=False)
    for idx, f in enumerate(ctx.findings, 1):
        evidence = f.evidence or "无"
        reason_points, suggestion_points = split_reason_suggestion(f)
        card_key = checklist_key(str(ctx.path), str(idx), f.item, f.location, f.message)
        checkable = f.status != "PASS"
        reason_html = render_check_items(card_key, "reason", reason_points, checkable)
        suggestion_html = render_check_items(card_key, "suggestion", suggestion_points, checkable)
        done_html = ""
        main_class = "finding-main"
        if checkable:
            done_html = f"""
                <label class="done-row">
                  <input type="checkbox" class="card-check" data-check-key="{card_key}:done">
                  <span>本项已处理</span>
                  <small class="done-progress">0/0 细节</small>
                </label>
            """
        else:
            main_class = "finding-main no-check"
        tech_html = "" if not evidence or evidence == "无" else f"""
                  <details class="tech-detail">
                    <summary>技术详情（可忽略）</summary>
                    <div class="tech-body">{evidence}</div>
                  </details>
                """
        cards.append(
            f"""
            <article class="finding {f.status.lower()}" data-status="{f.status}" data-check-card="{card_key}" data-text="{html.escape((f.item + ' ' + f.location + ' ' + f.message + ' ' + evidence).lower())}">
              <div class="{main_class}">
                <div class="status-pill {f.status.lower()}">{status_label.get(f.status, f.status)}</div>
                <div class="finding-body">
                  <div class="finding-topline">
                    <h2>{idx}. {html.escape(f.item)}</h2>
                    <span class="severity">{html.escape(f.severity)}</span>
                    <span class="location">{html.escape(f.location)}</span>
                  </div>
                  <p class="message">{html.escape(f.message)}</p>
                </div>
                {done_html}
                <button class="toggle" type="button" aria-expanded="false">详情</button>
              </div>
              <div class="evidence" hidden>
                <div class="human-detail">
                  <div class="detail-grid">
                    <section class="detail-panel">
                      <div class="detail-title">为什么报错</div>
                      {reason_html}
                    </section>
                    <section class="detail-panel">
                      <div class="detail-title">修改建议</div>
                      {suggestion_html}
                    </section>
                  </div>
                </div>
                {tech_html}
              </div>
            </article>
            """
        )
    generated = __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    html_doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>论文格式检测报告</title>
<style>
:root {{
  --ink:#172033;
  --muted:#64748b;
  --line:#dbe3ee;
  --paper:#ffffff;
  --bg:#f4f7fb;
  --pass:#128044;
  --pass-bg:#dcfce7;
  --warn:#a16207;
  --warn-bg:#fef3c7;
  --fail:#c62828;
  --fail-bg:#fee2e2;
}}
* {{ box-sizing:border-box; }}
body {{
  margin:0;
  color:var(--ink);
  background:
    radial-gradient(circle at top left, rgba(59,130,246,.12), transparent 34rem),
    linear-gradient(180deg, #f8fbff 0%, var(--bg) 26rem);
  font-family:"Microsoft YaHei", "PingFang SC", Arial, sans-serif;
  line-height:1.6;
}}
.shell {{ max-width:1180px; margin:0 auto; padding:30px 22px 42px; }}
.hero {{
  display:flex;
  justify-content:space-between;
  gap:22px;
  align-items:flex-end;
  margin-bottom:18px;
}}
h1 {{ margin:0 0 8px; font-size:30px; letter-spacing:0; }}
.meta {{ color:var(--muted); font-size:14px; overflow-wrap:anywhere; }}
.summary {{ display:grid; grid-template-columns:repeat(4, minmax(140px,1fr)); gap:12px; margin:18px 0; }}
.stat {{
  background:rgba(255,255,255,.88);
  border:1px solid var(--line);
  border-radius:8px;
  padding:14px 16px;
  box-shadow:0 10px 28px rgba(15,23,42,.06);
}}
.stat .num {{ display:block; font-size:30px; font-weight:800; line-height:1; }}
.stat .label {{ color:var(--muted); font-size:14px; }}
.stat.pass .num {{ color:var(--pass); }}
.stat.warn .num {{ color:var(--warn); }}
.stat.fail .num {{ color:var(--fail); }}
.stat.progress .num {{ color:#1d4ed8; }}
.progress-track {{
  height:7px;
  margin-top:10px;
  border-radius:999px;
  background:#e2e8f0;
  overflow:hidden;
}}
.progress-fill {{
  width:0%;
  height:100%;
  border-radius:999px;
  background:linear-gradient(90deg, #2563eb, #0f766e);
  transition:width .18s ease;
}}
.toolbar {{
  position:sticky;
  top:0;
  z-index:10;
  display:flex;
  align-items:center;
  gap:10px;
  flex-wrap:wrap;
  padding:12px;
  margin:18px 0;
  border:1px solid var(--line);
  border-radius:8px;
  background:rgba(255,255,255,.94);
  backdrop-filter:blur(10px);
  box-shadow:0 8px 20px rgba(15,23,42,.05);
}}
.toolbar-actions {{
  display:flex;
  gap:8px;
  flex-wrap:wrap;
}}
.filter-group {{
  display:flex;
  gap:8px;
  flex-wrap:wrap;
  align-items:center;
}}
.group-label {{
  color:var(--muted);
  font-size:13px;
  font-weight:700;
}}
.filter, .search {{
  height:36px;
  border:1px solid var(--line);
  background:#fff;
  border-radius:8px;
  color:var(--ink);
  font-size:14px;
}}
.filter {{ padding:0 12px; cursor:pointer; }}
.filter.active {{ border-color:#2563eb; color:#1d4ed8; background:#eff6ff; font-weight:700; }}
.search {{ flex:1; min-width:260px; padding:0 12px; }}
.action-btn {{
  height:36px;
  border:1px solid #bfd0e6;
  background:#f8fbff;
  border-radius:8px;
  color:#18345f;
  font-size:14px;
  font-weight:700;
  padding:0 12px;
  cursor:pointer;
}}
.action-btn.primary {{
  border-color:#1d4ed8;
  background:#2563eb;
  color:#fff;
}}
.action-btn:hover {{
  transform:translateY(-1px);
  box-shadow:0 6px 14px rgba(15,23,42,.08);
}}
.list {{ display:flex; flex-direction:column; gap:10px; }}
.finding {{
  border:1px solid var(--line);
  border-left-width:5px;
  border-radius:8px;
  background:var(--paper);
  overflow:hidden;
  box-shadow:0 8px 22px rgba(15,23,42,.045);
}}
.finding.checked {{
  background:#f9fbff;
  border-color:#cbd8ea;
}}
.finding.checked .message {{
  color:#64748b;
}}
.finding.pass {{ border-left-color:var(--pass); }}
.finding.warn {{ border-left-color:var(--warn); }}
.finding.fail {{ border-left-color:var(--fail); }}
.finding-main {{
  display:grid;
  grid-template-columns:92px minmax(0,1fr) 132px 74px;
  gap:14px;
  align-items:start;
  padding:14px 14px 14px 12px;
}}
.finding-main.no-check {{
  grid-template-columns:92px minmax(0,1fr) 74px;
}}
.status-pill {{
  width:76px;
  text-align:center;
  padding:5px 0;
  border-radius:999px;
  font-weight:800;
  font-size:13px;
}}
.status-pill.pass {{ color:var(--pass); background:var(--pass-bg); }}
.status-pill.warn {{ color:var(--warn); background:var(--warn-bg); }}
.status-pill.fail {{ color:var(--fail); background:var(--fail-bg); }}
.finding-topline {{ display:flex; align-items:center; gap:8px; flex-wrap:wrap; }}
.finding h2 {{ margin:0; font-size:17px; line-height:1.35; }}
.severity, .location {{
  display:inline-flex;
  align-items:center;
  min-height:24px;
  padding:2px 8px;
  border-radius:999px;
  font-size:12px;
  color:#475569;
  background:#f1f5f9;
}}
.message {{ margin:7px 0 0; color:#334155; overflow-wrap:anywhere; }}
.toggle {{
  border:1px solid var(--line);
  background:#f8fafc;
  border-radius:8px;
  padding:7px 10px;
  cursor:pointer;
  color:#0f172a;
}}
.toggle:hover {{ background:#eef2f7; }}
.done-row {{
  display:grid;
  grid-template-columns:18px 1fr;
  gap:4px 7px;
  align-items:center;
  min-height:36px;
  padding:7px 9px;
  border:1px solid #dbe7f5;
  border-radius:8px;
  background:#f8fafc;
  color:#334155;
  font-size:13px;
  font-weight:700;
  cursor:pointer;
  user-select:none;
}}
.done-row input, .check-row input {{
  width:16px;
  height:16px;
  accent-color:#2563eb;
  cursor:pointer;
}}
.done-row.checked {{
  border-color:#93c5fd;
  background:#eff6ff;
  color:#1d4ed8;
}}
.done-progress {{
  grid-column:2;
  color:#64748b;
  font-size:12px;
  font-weight:500;
  line-height:1.2;
}}
.check-list {{
  list-style:none;
  padding-left:0;
}}
.check-item {{
  display:flex;
  align-items:flex-start;
  gap:8px;
  margin:6px 0;
}}
.check-row {{
  display:grid;
  grid-template-columns:18px 1fr;
  gap:8px;
  align-items:start;
  flex:1;
  min-width:0;
  padding:6px 8px;
  border-radius:8px;
  cursor:pointer;
}}
.check-row:hover {{
  background:#f8fafc;
}}
.check-row.checked {{
  background:#eff6ff;
  color:#475569;
}}
.check-row.checked span {{
  text-decoration:line-through;
  text-decoration-thickness:1px;
  text-decoration-color:#94a3b8;
}}
.check-text {{
  overflow-wrap:anywhere;
}}
.copy-snippet {{
  flex:0 0 auto;
  border:1px solid #bfd0e6;
  background:#f8fbff;
  color:#1d4ed8;
  border-radius:8px;
  padding:5px 8px;
  margin-top:4px;
  font-size:12px;
  font-weight:700;
  cursor:pointer;
}}
.copy-snippet:hover {{
  background:#eff6ff;
  border-color:#93c5fd;
}}
.copy-snippet.copied {{
  color:#128044;
  border-color:#86efac;
  background:#dcfce7;
}}
.evidence {{
  border-top:1px solid #eef2f7;
  padding:12px 16px 14px 112px;
  color:#1f2937;
  background:#fbfdff;
  max-height:none;
  overflow:auto;
  overflow-wrap:anywhere;
  white-space:normal;
  font-family:"Microsoft YaHei", Arial, sans-serif;
  font-size:14px;
}}
.detail-title {{
  font-weight:800;
  margin-bottom:6px;
  color:#0f172a;
}}
.detail-grid {{
  display:grid;
  grid-template-columns:1fr 1fr;
  gap:12px;
}}
.detail-panel {{
  border:1px solid #e5edf7;
  border-radius:8px;
  background:#fff;
  padding:10px 12px;
}}
.plain-list {{
  margin:0;
  padding-left:20px;
}}
.plain-list li {{
  margin:3px 0;
}}
.tech-detail {{
  margin-top:12px;
  border-top:1px dashed #dbe3ee;
  padding-top:10px;
  color:#64748b;
}}
.tech-detail summary {{
  cursor:pointer;
  user-select:none;
  font-size:13px;
  color:#64748b;
}}
.tech-body {{
  margin-top:8px;
  max-height:170px;
  overflow:auto;
  padding:10px;
  border-radius:8px;
  background:#f8fafc;
  font-size:12px;
  color:#475569;
  overflow-wrap:anywhere;
}}
.empty {{ display:none; padding:30px; text-align:center; color:var(--muted); }}
@media (max-width:760px) {{
  .hero {{ display:block; }}
  .summary {{ grid-template-columns:1fr; }}
  .finding-main {{ grid-template-columns:1fr; }}
  .evidence {{ padding:12px 14px; }}
  .detail-grid {{ grid-template-columns:1fr; }}
  .toggle {{ width:100%; }}
  .done-row {{ width:100%; }}
}}
</style>
</head>
<body>
<main class="shell">
<section class="hero">
  <div>
    <h1>本科毕业设计(论文)格式检测报告</h1>
    <div class="meta">检测文件：{html.escape(str(ctx.path))}</div>
    <div class="meta">生成时间：{generated}</div>
  </div>
</section>
<section class="summary">
  <div class="stat pass"><span class="num">{counts.get('PASS',0)}</span><span class="label">通过</span></div>
  <div class="stat warn"><span class="num">{counts.get('WARN',0)}</span><span class="label">提醒</span></div>
  <div class="stat fail"><span class="num">{counts.get('FAIL',0)}</span><span class="label">未通过</span></div>
  <div class="stat progress">
    <span class="num" id="overall-progress">0%</span>
    <span class="label" id="overall-progress-text">待处理项加载中</span>
    <div class="progress-track" aria-hidden="true"><div class="progress-fill" id="overall-progress-fill"></div></div>
  </div>
</section>
<section class="toolbar">
  <div class="filter-group" aria-label="检测状态筛选">
    <span class="group-label">状态</span>
    <button class="filter status-filter active" data-filter="PROBLEM" type="button">待处理问题</button>
    <button class="filter status-filter" data-filter="FAIL" type="button">未通过</button>
    <button class="filter status-filter" data-filter="WARN" type="button">提醒</button>
    <button class="filter status-filter" data-filter="PASS" type="button">通过</button>
    <button class="filter status-filter" data-filter="ALL" type="button">全部</button>
  </div>
  <div class="filter-group" aria-label="处理状态筛选">
    <span class="group-label">处理</span>
    <button class="filter work-filter active" data-work-filter="ALL" type="button">全部</button>
    <button class="filter work-filter" data-work-filter="OPEN" type="button">未处理</button>
    <button class="filter work-filter" data-work-filter="DONE" type="button">已处理</button>
  </div>
  <input class="search" type="search" placeholder="搜索检测项、位置、说明或证据">
  <div class="toolbar-actions">
    <button class="action-btn" id="copy-open" type="button">复制未处理清单</button>
    <button class="action-btn primary" id="export-state" type="button">导出带勾选状态的HTML</button>
    <button class="action-btn" id="clear-state" type="button">清空勾选</button>
  </div>
</section>
<section class="list">
  {''.join(cards)}
  <div class="empty">没有匹配的检测项。</div>
</section>
</main>
<script id="embedded-check-state" type="application/json">{embedded_state_json}</script>
<script>
const REPORT_KEY = '{report_key}';
const STORAGE_KEY = 'thesis-format-audit:' + REPORT_KEY;
const statusFilters = [...document.querySelectorAll('.status-filter')];
const workFilters = [...document.querySelectorAll('.work-filter')];
const search = document.querySelector('.search');
const findings = [...document.querySelectorAll('.finding')];
const empty = document.querySelector('.empty');
const progressNum = document.getElementById('overall-progress');
const progressText = document.getElementById('overall-progress-text');
const progressFill = document.getElementById('overall-progress-fill');
let activeStatus = 'PROBLEM';
let activeWork = 'ALL';
let checkState = loadCheckState();

function applyFilter() {{
  const q = search.value.trim().toLowerCase();
  let shown = 0;
  findings.forEach(card => {{
    const isProblem = card.dataset.status !== 'PASS';
    const isDone = Boolean(card.querySelector('.card-check')?.checked);
    const statusOk =
      activeStatus === 'ALL' ||
      (activeStatus === 'PROBLEM' && isProblem) ||
      card.dataset.status === activeStatus;
    const workOk =
      activeWork === 'ALL' ||
      (activeWork === 'OPEN' && isProblem && !isDone) ||
      (activeWork === 'DONE' && isProblem && isDone);
    const textOk = !q || card.dataset.text.includes(q);
    const visible = statusOk && workOk && textOk;
    card.style.display = visible ? '' : 'none';
    if (visible) shown += 1;
  }});
  empty.style.display = shown ? 'none' : 'block';
}}

statusFilters.forEach(btn => btn.addEventListener('click', () => {{
  activeStatus = btn.dataset.filter;
  statusFilters.forEach(x => x.classList.toggle('active', x === btn));
  applyFilter();
}}));

workFilters.forEach(btn => btn.addEventListener('click', () => {{
  activeWork = btn.dataset.workFilter;
  workFilters.forEach(x => x.classList.toggle('active', x === btn));
  applyFilter();
}}));
search.addEventListener('input', applyFilter);

document.querySelectorAll('.toggle').forEach(btn => btn.addEventListener('click', () => {{
  const card = btn.closest('.finding');
  const evidence = card.querySelector('.evidence');
  const open = evidence.hasAttribute('hidden');
  if (open) {{
    evidence.removeAttribute('hidden');
    btn.textContent = '收起';
    btn.setAttribute('aria-expanded', 'true');
  }} else {{
    evidence.setAttribute('hidden', '');
    btn.textContent = '详情';
    btn.setAttribute('aria-expanded', 'false');
  }}
}}));

function parseJsonSafe(text) {{
  try {{
    return JSON.parse(text || '{{}}');
  }} catch (err) {{
    return {{}};
  }}
}}

function loadCheckState() {{
  const embedded = parseJsonSafe(document.getElementById('embedded-check-state')?.textContent || '{{}}');
  const stored = parseJsonSafe(readStoredState());
  return {{...embedded, ...stored}};
}}

function readStoredState() {{
  try {{
    return localStorage.getItem(STORAGE_KEY) || '{{}}';
  }} catch (err) {{
    return '{{}}';
  }}
}}

function saveCheckState() {{
  try {{
    localStorage.setItem(STORAGE_KEY, JSON.stringify(checkState));
  }} catch (err) {{
    // 个别浏览器会限制本地 HTML 使用 localStorage；导出按钮仍会把勾选状态写入 HTML。
  }}
}}

function updateCardProgress(card) {{
  const subChecks = [...card.querySelectorAll('.sub-check')];
  const done = card.querySelector('.card-check');
  const progress = card.querySelector('.done-progress');
  const checkedCount = subChecks.filter(x => x.checked).length;
  if (progress) progress.textContent = `${{checkedCount}}/${{subChecks.length}} 细节`;
  card.classList.toggle('checked', Boolean(done?.checked));
  const doneRow = done?.closest('.done-row');
  if (doneRow) doneRow.classList.toggle('checked', Boolean(done?.checked));
}}

function updateOverallProgress() {{
  const problemChecks = [...document.querySelectorAll('.card-check')];
  const total = problemChecks.length;
  const done = problemChecks.filter(x => x.checked).length;
  const percent = total ? Math.round(done * 100 / total) : 100;
  progressNum.textContent = `${{percent}}%`;
  progressText.textContent = total ? `已处理 ${{done}} / ${{total}} 项` : '没有待处理项';
  progressFill.style.width = `${{percent}}%`;
}}

function applyCheckVisual(input) {{
  const row = input.closest('.check-row') || input.closest('.done-row');
  if (row) row.classList.toggle('checked', input.checked);
  const card = input.closest('.finding');
  if (card) updateCardProgress(card);
}}

function syncAllChecks() {{
  document.querySelectorAll('[data-check-key]').forEach(input => {{
    input.checked = Boolean(checkState[input.dataset.checkKey]);
    applyCheckVisual(input);
  }});
  updateOverallProgress();
  applyFilter();
}}

document.querySelectorAll('[data-check-key]').forEach(input => {{
  input.addEventListener('change', () => {{
    if (input.checked) {{
      checkState[input.dataset.checkKey] = true;
    }} else {{
      delete checkState[input.dataset.checkKey];
    }}
    saveCheckState();
    applyCheckVisual(input);
    updateOverallProgress();
    applyFilter();
  }});
}});

function currentStateFromDom() {{
  const state = {{}};
  document.querySelectorAll('[data-check-key]').forEach(input => {{
    if (input.checked) state[input.dataset.checkKey] = true;
  }});
  return state;
}}

document.getElementById('export-state').addEventListener('click', () => {{
  checkState = currentStateFromDom();
  saveCheckState();
  const clone = document.documentElement.cloneNode(true);
  const stateTag = clone.querySelector('#embedded-check-state');
  if (stateTag) stateTag.textContent = JSON.stringify(checkState, null, 2);
  const htmlText = '<!doctype html>\\n' + clone.outerHTML;
  const blob = new Blob([htmlText], {{type:'text/html;charset=utf-8'}});
  const a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = '论文格式检测报告_含勾选记录.html';
  document.body.appendChild(a);
  a.click();
  a.remove();
  setTimeout(() => URL.revokeObjectURL(a.href), 800);
}});

function openProblemCards() {{
  return findings.filter(card => card.dataset.status !== 'PASS' && !card.querySelector('.card-check')?.checked);
}}

function cardSummary(card, idx) {{
  const status = card.querySelector('.status-pill')?.textContent.trim() || '';
  const title = card.querySelector('h2')?.textContent.trim() || '未命名检测项';
  const location = card.querySelector('.location')?.textContent.trim() || '';
  const message = card.querySelector('.message')?.textContent.trim() || '';
  const panels = [...card.querySelectorAll('.detail-panel')];
  const suggestions = panels[1] ? [...panels[1].querySelectorAll('li')].map(x => x.textContent.trim()).filter(Boolean) : [];
  const suggestionText = suggestions.length ? `\\n   修改建议：${{suggestions.join('；')}}` : '';
  return `${{idx}}. [${{status}}] ${{title}}\\n   位置：${{location}}\\n   问题：${{message}}${{suggestionText}}`;
}}

async function copyText(text) {{
  try {{
    await navigator.clipboard.writeText(text);
    return true;
  }} catch (err) {{
    const box = document.createElement('textarea');
    box.value = text;
    box.style.position = 'fixed';
    box.style.left = '-9999px';
    document.body.appendChild(box);
    box.select();
    const ok = document.execCommand('copy');
    box.remove();
    return ok;
  }}
}}

document.querySelectorAll('.copy-snippet').forEach(btn => {{
  btn.addEventListener('click', async (event) => {{
    event.preventDefault();
    event.stopPropagation();
    const ok = await copyText(btn.dataset.copyText || '');
    const oldText = btn.textContent;
    btn.textContent = ok ? '已复制' : '复制失败';
    btn.classList.toggle('copied', ok);
    setTimeout(() => {{
      btn.textContent = oldText;
      btn.classList.remove('copied');
    }}, 1200);
  }});
}});

document.getElementById('copy-open').addEventListener('click', async () => {{
  const openCards = openProblemCards();
  const text = openCards.length
    ? '未处理问题清单：\\n\\n' + openCards.map((card, idx) => cardSummary(card, idx + 1)).join('\\n\\n')
    : '未处理问题清单：\\n\\n当前没有未处理的问题项。';
  const ok = await copyText(text);
  alert(ok ? `已复制 ${{openCards.length}} 个未处理问题。` : '复制失败，请手动选择页面内容复制。');
}});

document.getElementById('clear-state').addEventListener('click', () => {{
  if (!confirm('确定要清空本报告中的所有勾选记录吗？')) return;
  checkState = {{}};
  saveCheckState();
  syncAllChecks();
}});

syncAllChecks();
</script>
</body>
</html>"""
    out.write_text(html_doc, encoding="utf-8")


def run_audit(path: Path, out: Path) -> AuditContext:
    with tempfile.TemporaryDirectory(prefix="docx-compat-") as tmp:
        doc, audit_path = open_docx_document(path, Path(tmp))
        ctx = AuditContext(path=audit_path, doc=doc)
        audit_basic(ctx)
        audit_template_residue(ctx)
        audit_abstracts(ctx)
        audit_toc(ctx)
        audit_headings(ctx)
        audit_body_and_numbering(ctx)
        audit_chinese_punctuation(ctx)
        audit_figures_tables(ctx)
        audit_tables(ctx)
        audit_units(ctx)
        audit_formulas(ctx)
        audit_references(ctx)
        audit_document_structure(ctx)
        audit_word_count_requirements(ctx)
        audit_appendix(ctx)
        audit_acknowledgement(ctx)
        audit_headers_footers_structural(ctx)
        audit_page_numbers(ctx)
        audit_vip_compatibility_summary(ctx)
        build_report(ctx, out)
        return ctx


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="本科毕业设计论文 Word 格式检测脚本")
    parser.add_argument("docx", help="要检测的 .docx 文件路径")
    parser.add_argument("--out", help="HTML 报告输出路径。默认在 docx 同目录生成 *_format_report.html")
    args = parser.parse_args(argv)

    path = Path(args.docx).expanduser().resolve()
    if not path.exists():
        print(f"文件不存在：{path}", file=sys.stderr)
        return 2
    if path.suffix.lower() != ".docx":
        print("当前脚本只支持 .docx 文件。请先将 .doc 转为 .docx。", file=sys.stderr)
        return 2
    out = Path(args.out).expanduser().resolve() if args.out else path.with_name(path.stem + "_format_report.html")
    ctx = run_audit(path, out)
    counts = Counter(f.status for f in ctx.findings)
    print(f"检测完成：{path}")
    print(f"报告输出：{out}")
    print(f"PASS={counts.get('PASS',0)} WARN={counts.get('WARN',0)} FAIL={counts.get('FAIL',0)}")
    print("\n未通过/提醒项：")
    for f in ctx.findings:
        if f.status != "PASS":
            print(f"[{f.status}] {f.item} - {f.location}: {f.message} {f.evidence}")
    return 1 if counts.get("FAIL", 0) else 0


if __name__ == "__main__":
    raise SystemExit(main())
